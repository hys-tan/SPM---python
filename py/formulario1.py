import tkinter as tk
from tkinter import PhotoImage
from tkinter import ttk, messagebox
from tkinter import filedialog
from tkcalendar import DateEntry
from docx import Document
from num2words import num2words
import os
import shutil
import utils


# Ruta base dinámica (directorio del archivo principal)
BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))

# Ruta a la carpeta de iconos
ICON_DIR = os.path.join(BASE_DIR, "icons")

# RUTA PARA WORD PLANTILLA
COT_PLAN = os.path.join(BASE_DIR, "plantilla")

# NOTIFICACIONES - CONFIRMACIONES
class alertas:
    def __init__(self, parent):
        self.parent = parent
    
    # CERRAR EL PROGRAMA --
    def cerrar_prog(self):
        salida=tk.Toplevel(self.parent)
        salida.title("")
        salida.geometry("300x110")
        salida.resizable(False, False)
        salida.configure(bg="#FFFFFF")
        salida.grab_set()
        utils.centrar_ventana(salida)
        salida.protocol("WM_DELETE_WINDOW", lambda: None)
        
        canvas_salida = tk.Canvas(salida, width=300, height=110, bg="#FFFFFF", highlightthickness=0)
        canvas_salida.pack()
        
        icono_path = os.path.join(ICON_DIR, "question.png")
        try:
            icono_alert_ex = tk.PhotoImage(file=icono_path)
            canvas_salida.create_image(30, 17, anchor="nw", image=icono_alert_ex)
            canvas_salida.image = icono_alert_ex
        except Exception as e:
            raise FileNotFoundError(f"El archivo del icono no se encontró en la ruta: {icono_path}. Error: {e}")

        utils.create_rounded_rectangle(canvas_salida, 0, 66, 300, 110, radius=0, fill="#EEEEE4", outline="#EEEEE4")
        
        canvas_salida.create_text(79, 26, text="¿Está seguro de que desea salir?", anchor="nw", font=("Arial", 10), fill="Black")
        
        btn_si = tk.Button(salida, text="Si", width=9, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=self.parent.destroy)
        btn_si.place(x=73, y=73)
                
        btn_no = tk.Button(salida, text="No", width=9, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=salida.destroy)
        btn_no.place(x=158, y=73)
        
        utils.aplicar_hover_a_botones([btn_si, btn_no])

    # SELECCIONE UNA FILA
    def seleccionar_fila(self):
        sec_fila=tk.Toplevel(self.parent)
        sec_fila.title("")
        sec_fila.geometry("300x110")
        sec_fila.resizable(False, False)
        sec_fila.configure(bg="#FFFFFF")
        sec_fila.grab_set()
        utils.centrar_ventana(sec_fila)
        sec_fila.protocol("WM_DELETE_WINDOW", lambda: None)
        
        canvas_sec_fila = tk.Canvas(sec_fila, width=300, height=110, bg="#FFFFFF", highlightthickness=0)
        canvas_sec_fila.pack()
        
        # Cargar el icono
        icono_path = os.path.join(ICON_DIR, "alert.png")
        try:
            icono_alert_sec = tk.PhotoImage(file=icono_path)
            canvas_sec_fila.create_image(30, 17, anchor="nw", image=icono_alert_sec)
            canvas_sec_fila.image = icono_alert_sec
        except Exception as e:
            raise FileNotFoundError(f"El archivo del icono no se encontró en la ruta: {icono_path}. Error: {e}")

        utils.create_rounded_rectangle(canvas_sec_fila, 0, 66, 300, 110, radius=0, fill="#EEEEE4", outline="#EEEEE4")
        
        canvas_sec_fila.create_text(79, 26, text="Por favor, seleccione una fila", anchor="nw", font=("Arial", 10), fill="Black")
        
        btn_sec_ok = tk.Button(sec_fila, text="Aceptar", width=9, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=sec_fila.destroy)
        btn_sec_ok.place(x=115, y=73)
        
        utils.aplicar_hover_a_botones([btn_sec_ok])

    # DATO NO ENCONTRADO
    def no_datos(self):
        no_file=tk.Toplevel(self.parent)
        no_file.title("")
        no_file.geometry("300x110")
        no_file.resizable(False, False)
        no_file.configure(bg="#FFFFFF")
        no_file.grab_set()
        utils.centrar_ventana(no_file)
        no_file.protocol("WM_DELETE_WINDOW", lambda: None)
        
        canvas_no_file = tk.Canvas(no_file, width=300, height=110, bg="#FFFFFF", highlightthickness=0)
        canvas_no_file.pack()
        
        # Cargar el icono
        icono_path = os.path.join(ICON_DIR, "delete.png")
        try:
            icono_alert_nf = tk.PhotoImage(file=icono_path)
            canvas_no_file.create_image(30, 17, anchor="nw", image=icono_alert_nf)
            canvas_no_file.image = icono_alert_nf
        except Exception as e:
            raise FileNotFoundError(f"El archivo del icono no se encontró en la ruta: {icono_path}. Error: {e}")

        utils.create_rounded_rectangle(canvas_no_file, 0, 66, 300, 110, radius=0, fill="#EEEEE4", outline="#EEEEE4")
        
        canvas_no_file.create_text(94, 26, text="Datos no encontrados", anchor="nw", font=("Arial", 10), fill="Black")
        
        btn_nf_ok = tk.Button(no_file, text="Aceptar", width=9, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=no_file.destroy)
        btn_nf_ok.place(x=115, y=73)
        
        utils.aplicar_hover_a_botones([btn_nf_ok])

    # PREGUNTA ELIMINAR FILA
    def quest_mat(self):
        q_material=tk.Toplevel(self.parent)
        q_material.title("")
        q_material.geometry("300x110")
        q_material.resizable(False, False)
        q_material.configure(bg="#FFFFFF")
        q_material.grab_set()
        utils.centrar_ventana(q_material)
        q_material.protocol("WM_DELETE_WINDOW", lambda: None)
        
        canvas_ques_mat = tk.Canvas(q_material, width=300, height=110, bg="#FFFFFF", highlightthickness=0)
        canvas_ques_mat.pack()

        icono_path = os.path.join(ICON_DIR, "alert.png")
        try:
            icon_quest_mat = tk.PhotoImage(file=icono_path)
            canvas_ques_mat.create_image(30, 17, anchor="nw", image=icon_quest_mat)
            canvas_ques_mat.image = icon_quest_mat
        except Exception as e:
            raise FileNotFoundError(f"El archivo del icono no se encontró en la ruta: {icono_path}. Error: {e}")

        utils.create_rounded_rectangle(canvas_ques_mat, 0, 66, 300, 110, radius=0, fill="#EEEEE4", outline="#EEEEE4")
        
        canvas_ques_mat.create_text(79, 19, text="¿Estás seguro de que deseas", anchor="nw", font=("Arial", 10), fill="Black")
        canvas_ques_mat.create_text(115, 33, text="eliminar esta fila?", anchor="nw", font=("Arial", 10), fill="Black")
        
        btn_qm_si = tk.Button(q_material, text="Si", width=9, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_qm_si.place(x=73, y=73)

        btn_qm_no = tk.Button(q_material, text="No", width=9, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=q_material.destroy)
        btn_qm_no.place(x=158, y=73)
        
        utils.aplicar_hover_a_botones([btn_qm_si, btn_qm_no])

    # FILA ELIMINADA
    def material_delete(self):
        mat_delete=tk.Toplevel(self.parent)
        mat_delete.title("")
        mat_delete.geometry("300x110")
        mat_delete.resizable(False, False)
        mat_delete.configure(bg="#FFFFFF")
        mat_delete.grab_set()
        utils.centrar_ventana(mat_delete)
        mat_delete.protocol("WM_DELETE_WINDOW", lambda: None)
        
        canvas_mat_delete = tk.Canvas(mat_delete, width=300, height=110, bg="#FFFFFF", highlightthickness=0)
        canvas_mat_delete.pack()
        
        icono_path = os.path.join(ICON_DIR, "confirm.png")
        try:
            icon_confirm_del = tk.PhotoImage(file=icono_path)
            canvas_mat_delete.create_image(30, 17, anchor="nw", image=icon_confirm_del)
            canvas_mat_delete.image = icon_confirm_del
        except Exception as e:
            raise FileNotFoundError(f"El archivo del icono no se encontró en la ruta: {icono_path}. Error: {e}")
        
        utils.create_rounded_rectangle(canvas_mat_delete, 0, 66, 300, 110, radius=0, fill="#EEEEE4", outline="#EEEEE4")
        
        canvas_mat_delete.create_text(85, 26, text="Fila eliminada exitosamente", anchor="nw", font=("Arial", 10), fill="Black")
        
        btn_md_ok = tk.Button(mat_delete, text="Aceptar", width=9, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=mat_delete.destroy)
        btn_md_ok.place(x=115, y=73)
        
        utils.aplicar_hover_a_botones([btn_md_ok])
    
    # GENERAR COTIZACION
    def generator_cotizacion(self):
        gener_coti=tk.Toplevel(self.parent)
        gener_coti.title("")
        gener_coti.geometry("300x110")
        gener_coti.resizable(False, False)
        gener_coti.configure(bg="#FFFFFF")
        gener_coti.grab_set()
        utils.centrar_ventana(gener_coti)
        gener_coti.protocol("WM_DELETE_WINDOW", lambda: None)
        
        can_gen_coti = tk.Canvas(gener_coti, width=300, height=110, bg="#FFFFFF", highlightthickness=0)
        can_gen_coti.pack()
        
        icono_path = os.path.join(ICON_DIR, "question.png")
        try:
            icon_gener_coti = tk.PhotoImage(file=icono_path)
            can_gen_coti.create_image(30, 17, anchor="nw", image=icon_gener_coti)
            can_gen_coti.image = icon_gener_coti
        except Exception as e:
            raise FileNotFoundError(f"El archivo del icono no se encontró en la ruta: {icono_path}. Error: {e}")
        
        utils.create_rounded_rectangle(can_gen_coti, 0, 66, 300, 110, radius=0, fill="#EEEEE4", outline="#EEEEE4")
        
        can_gen_coti.create_text(79, 26, text="¿Desea generar la cotización?", anchor="nw", font=("Arial", 10), fill="Black")
        
        btn_gc_si = tk.Button(gener_coti, text="Si", width=9, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_gc_si.place(x=73, y=73)

        btn_gc_no = tk.Button(gener_coti, text="No", width=9, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=gener_coti.destroy)
        btn_gc_no.place(x=158, y=73)
        
        utils.aplicar_hover_a_botones([btn_gc_si, btn_gc_no])
    
    # CANCELAR COTIZACION
    def cancelar_cotizacion(self):
        cancel_coti=tk.Toplevel(self.parent)
        cancel_coti.title("")
        cancel_coti.geometry("300x110")
        cancel_coti.resizable(False, False)
        cancel_coti.configure(bg="#FFFFFF")
        cancel_coti.grab_set()
        utils.centrar_ventana(cancel_coti)
        cancel_coti.protocol("WM_DELETE_WINDOW", lambda: None)
        
        can_cancel_coti = tk.Canvas(cancel_coti, width=300, height=110, bg="#FFFFFF", highlightthickness=0)
        can_cancel_coti.pack()
        
        icono_path = os.path.join(ICON_DIR, "alert.png")
        try:
            icon_cancel_coti = tk.PhotoImage(file=icono_path)
            can_cancel_coti.create_image(30, 17, anchor="nw", image=icon_cancel_coti)
            can_cancel_coti.image = icon_cancel_coti
        except Exception as e:
            raise FileNotFoundError(f"El archivo del icono no se encontró en la ruta: {icono_path}. Error: {e}")
        
        utils.create_rounded_rectangle(can_cancel_coti, 0, 66, 300, 110, radius=0, fill="#EEEEE4", outline="#EEEEE4")
        
        can_cancel_coti.create_text(79, 26, text="¿Desea cancelar la cotización?", anchor="nw", font=("Arial", 10), fill="Black")
        
        btn_ccoti_si = tk.Button(cancel_coti, text="Si", width=9, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_ccoti_si.place(x=73, y=73)

        btn_ccoti_no = tk.Button(cancel_coti, text="No", width=9, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=cancel_coti.destroy)
        btn_ccoti_no.place(x=158, y=73)
        
        utils.aplicar_hover_a_botones([btn_ccoti_si, btn_ccoti_no])
    
    # CAMBIOS REALIZADOS
    def cambios_realizados(self):
        cambios_confirm=tk.Toplevel(self.parent)
        cambios_confirm.title("")
        cambios_confirm.geometry("300x110")
        cambios_confirm.resizable(False, False)
        cambios_confirm.configure(bg="#FFFFFF")
        cambios_confirm.grab_set()
        utils.centrar_ventana(cambios_confirm)
        cambios_confirm.protocol("WM_DELETE_WINDOW", lambda: None)
        
        canvas_camb_confirm = tk.Canvas(cambios_confirm, width=300, height=110, bg="#FFFFFF", highlightthickness=0)
        canvas_camb_confirm.pack()
        
        icono_path = os.path.join(ICON_DIR, "confirm.png")
        try:
            icon_confirm_cambio = tk.PhotoImage(file=icono_path)
            canvas_camb_confirm.create_image(30, 17, anchor="nw", image=icon_confirm_cambio)
            canvas_camb_confirm.image = icon_confirm_cambio
        except Exception as e:
            raise FileNotFoundError(f"El archivo del icono no se encontró en la ruta: {icono_path}. Error: {e}")
        
        utils.create_rounded_rectangle(canvas_camb_confirm, 0, 66, 300, 110, radius=0, fill="#EEEEE4", outline="#EEEEE4")
        
        canvas_camb_confirm.create_text(80, 26, text="Cambios realizados con éxtio", anchor="nw", font=("Arial", 10), fill="Black")

        btn_cr_ok = tk.Button(cambios_confirm, text="Aceptar", width=9, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=cambios_confirm.destroy)
        btn_cr_ok.place(x=115, y=73)
        
        utils.aplicar_hover_a_botones([btn_cr_ok])

    # CONFIRMACION DE CONTIZACION GENERADA
    def confirmacion_cotizacion(self):
        coti_confirm=tk.Toplevel(self.parent)
        coti_confirm.title("")
        coti_confirm.geometry("300x110")
        coti_confirm.resizable(False, False)
        coti_confirm.configure(bg="#FFFFFF")
        coti_confirm.grab_set()
        utils.centrar_ventana(coti_confirm)
        coti_confirm.protocol("WM_DELETE_WINDOW", lambda: None)
        
        canvas_coti_confirm = tk.Canvas(coti_confirm, width=300, height=110, bg="#FFFFFF", highlightthickness=0)
        canvas_coti_confirm.pack()
        
        icono_path = os.path.join(ICON_DIR, "confirm.png")
        try:
            icon_confirm_coti = tk.PhotoImage(file=icono_path)
            canvas_coti_confirm.create_image(30, 17, anchor="nw", image=icon_confirm_coti)
            canvas_coti_confirm.image = icon_confirm_coti
        except Exception as e:
            raise FileNotFoundError(f"El archivo del icono no se encontró en la ruta: {icono_path}. Error: {e}")
        
        utils.create_rounded_rectangle(canvas_coti_confirm, 0, 66, 300, 110, radius=0, fill="#EEEEE4", outline="#EEEEE4")
        
        canvas_coti_confirm.create_text(93, 26, text="Cotización generada", anchor="nw", font=("Arial", 10), fill="Black")
        
        btn_coti_gen = tk.Button(coti_confirm, text="Aceptar", width=9, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=coti_confirm.destroy)
        btn_coti_gen.place(x=115, y=73)
        
        utils.aplicar_hover_a_botones([btn_coti_gen])

    # COMPLETE LOS DATOS
    def question_datos(self):
        quest_datos=tk.Toplevel(self.parent)
        quest_datos.title("")
        quest_datos.geometry("300x110")
        quest_datos.resizable(False, False)
        quest_datos.configure(bg="#FFFFFF")
        quest_datos.grab_set()
        utils.centrar_ventana(quest_datos)
        quest_datos.protocol("WM_DELETE_WINDOW", lambda: None)
        
        canvas_quest_datos = tk.Canvas(quest_datos, width=300, height=110, bg="#FFFFFF", highlightthickness=0)
        canvas_quest_datos.pack()
        
        icono_path = os.path.join(ICON_DIR, "alert.png")
        try:
            icon_quest_datos = tk.PhotoImage(file=icono_path)
            canvas_quest_datos.create_image(30, 17, anchor="nw", image=icon_quest_datos)
            canvas_quest_datos.image = icon_quest_datos
        except Exception as e:
            raise FileNotFoundError(f"El archivo del icono no se encontró en la ruta: {icono_path}. Error: {e}")
        
        utils.create_rounded_rectangle(canvas_quest_datos, 0, 66, 300, 110, radius=0, fill="#EEEEE4", outline="#EEEEE4")
        
        canvas_quest_datos.create_text(84, 19, text="Por favor, complete todos", anchor="nw", font=("Arial", 10), fill="Black")
        canvas_quest_datos.create_text(125, 33, text="los datos", anchor="nw", font=("Arial", 10), fill="Black")
        
        btn_qd_ok = tk.Button(quest_datos, text="Aceptar", width=9, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=quest_datos.destroy)
        btn_qd_ok.place(x=115, y=73)
        
        utils.aplicar_hover_a_botones([btn_qd_ok])
        
    # REGISTRO CORRECTO
    def registro_confirm(self):
        reg_confirm=tk.Toplevel(self.parent)
        reg_confirm.title("")
        reg_confirm.geometry("300x110")
        reg_confirm.resizable(False, False)
        reg_confirm.configure(bg="#FFFFFF")
        reg_confirm.grab_set()
        utils.centrar_ventana(reg_confirm)
        reg_confirm.protocol("WM_DELETE_WINDOW", lambda: None)
        
        canvas_reg_confirm = tk.Canvas(reg_confirm, width=300, height=110, bg="#FFFFFF", highlightthickness=0)
        canvas_reg_confirm.pack()
        
        icono_path = os.path.join(ICON_DIR, "confirm.png")
        try:
            icon_reg_confirm = tk.PhotoImage(file=icono_path)
            canvas_reg_confirm.create_image(30, 17, anchor="nw", image=icon_reg_confirm)
            canvas_reg_confirm.image = icon_reg_confirm
        except Exception as e:
            raise FileNotFoundError(f"El archivo del icono no se encontró en la ruta: {icono_path}. Error: {e}")
        
        utils.create_rounded_rectangle(canvas_reg_confirm, 0, 66, 300, 110, radius=0, fill="#EEEEE4", outline="#EEEEE4")
        
        canvas_reg_confirm.create_text(103, 26, text="¡Registro exitoso!", anchor="nw", font=("Arial", 10), fill="Black")
        
        btn_reg_conf = tk.Button(reg_confirm, text="Aceptar", width=9, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=reg_confirm.destroy)
        btn_reg_conf.place(x=115, y=73)
        
        utils.aplicar_hover_a_botones([btn_reg_conf])
        
    # CANCELAR REGISTRO
    def cancelar_registro(self):
        cancel_reg=tk.Toplevel(self.parent)
        cancel_reg.title("")
        cancel_reg.geometry("300x110")
        cancel_reg.resizable(False, False)
        cancel_reg.configure(bg="#FFFFFF")
        cancel_reg.grab_set()
        utils.centrar_ventana(cancel_reg)
        cancel_reg.protocol("WM_DELETE_WINDOW", lambda: None)
        
        can_cancel_reg = tk.Canvas(cancel_reg, width=300, height=110, bg="#FFFFFF", highlightthickness=0)
        can_cancel_reg.pack()
        
        icono_path = os.path.join(ICON_DIR, "alert.png")
        try:
            icon_cancel_reg = tk.PhotoImage(file=icono_path)
            can_cancel_reg.create_image(30, 17, anchor="nw", image=icon_cancel_reg)
            can_cancel_reg.image = icon_cancel_reg
        except Exception as e:
            raise FileNotFoundError(f"El archivo del icono no se encontró en la ruta: {icono_path}. Error: {e}")
        
        utils.create_rounded_rectangle(can_cancel_reg, 0, 66, 300, 110, radius=0, fill="#EEEEE4", outline="#EEEEE4")
        
        can_cancel_reg.create_text(76, 26, text="¿Desea cancelar el registro?", anchor="nw", font=("Arial", 10), fill="Black")
        
        btn_cc_si = tk.Button(cancel_reg, text="Si", width=9, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_cc_si.place(x=73, y=73)

        btn_cc_no = tk.Button(cancel_reg, text="No", width=9, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=cancel_reg.destroy)
        btn_cc_no.place(x=158, y=73)
        
        utils.aplicar_hover_a_botones([btn_cc_si, btn_cc_no])
        
    # CAMBIOS SIN GUARDAR
    def cambios_sin_guardar(self):
        unsaved_changes=tk.Toplevel(self.parent)
        unsaved_changes.title("")
        unsaved_changes.geometry("300x110")
        unsaved_changes.resizable(False, False)
        unsaved_changes.configure(bg="#FFFFFF")
        unsaved_changes.grab_set()
        utils.centrar_ventana(unsaved_changes)
        unsaved_changes.protocol("WM_DELETE_WINDOW", lambda: None)
        
        can_uns_ch = tk.Canvas(unsaved_changes, width=300, height=110, bg="#FFFFFF", highlightthickness=0)
        can_uns_ch.pack()
        
        icono_path = os.path.join(ICON_DIR, "alert.png")
        try:
            icon_uns_ch = tk.PhotoImage(file=icono_path)
            can_uns_ch.create_image(30, 17, anchor="nw", image=icon_uns_ch)
            can_uns_ch.image = icon_uns_ch
        except Exception as e:
            raise FileNotFoundError(f"El archivo del icono no se encontró en la ruta: {icono_path}. Error: {e}")
        
        utils.create_rounded_rectangle(can_uns_ch, 0, 66, 300, 110, radius=0, fill="#EEEEE4", outline="#EEEEE4")
        
        can_uns_ch.create_text(79, 19, text="Tiene cambios sin guardar,", anchor="nw", font=("Arial", 10), fill="Black")
        can_uns_ch.create_text(102, 33, text="¿Desea cancelar?", anchor="nw", font=("Arial", 10), fill="Black")
        
        btn_uns_si = tk.Button(unsaved_changes, text="Si", width=9, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_uns_si.place(x=73, y=73)

        btn_uns_no = tk.Button(unsaved_changes, text="No", width=9, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=unsaved_changes.destroy)
        btn_uns_no.place(x=158, y=73)
        
        utils.aplicar_hover_a_botones([btn_uns_si, btn_uns_no])
        
    # REGISTRO ACTUALIZADO
    def registro_actualizado(self):
        reg_updated=tk.Toplevel(self.parent)
        reg_updated.title("")
        reg_updated.geometry("300x110")
        reg_updated.resizable(False, False)
        reg_updated.configure(bg="#FFFFFF")
        reg_updated.grab_set()
        utils.centrar_ventana(reg_updated)
        reg_updated.protocol("WM_DELETE_WINDOW", lambda: None)
        
        canvas_reg_updated = tk.Canvas(reg_updated, width=300, height=110, bg="#FFFFFF", highlightthickness=0)
        canvas_reg_updated.pack()
        
        icono_path = os.path.join(ICON_DIR, "confirm.png")
        try:
            icon_reg_updated = tk.PhotoImage(file=icono_path)
            canvas_reg_updated.create_image(30, 17, anchor="nw", image=icon_reg_updated)
            canvas_reg_updated.image = icon_reg_updated
        except Exception as e:
            raise FileNotFoundError(f"El archivo del icono no se encontró en la ruta: {icono_path}. Error: {e}")
        
        utils.create_rounded_rectangle(canvas_reg_updated, 0, 66, 300, 110, radius=0, fill="#EEEEE4", outline="#EEEEE4")
        
        canvas_reg_updated.create_text(92, 26, text="¡Registro actualizado!", anchor="nw", font=("Arial", 10), fill="Black")
        
        btn_reg_upd = tk.Button(reg_updated, text="Aceptar", width=9, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=reg_updated.destroy)
        btn_reg_upd.place(x=115, y=73)
        
        utils.aplicar_hover_a_botones([btn_reg_upd])
    
    # CODIGO YA UTILIZADO
    def codigo_utilizado(self):
        code_used=tk.Toplevel(self.parent)
        code_used.title("")
        code_used.geometry("300x110")
        code_used.resizable(False, False)
        code_used.configure(bg="#FFFFFF")
        code_used.grab_set()
        utils.centrar_ventana(code_used)
        code_used.protocol("WM_DELETE_WINDOW", lambda: None)
        
        canvas_code_used = tk.Canvas(code_used, width=300, height=110, bg="#FFFFFF", highlightthickness=0)
        canvas_code_used.pack()
        
        icono_path = os.path.join(ICON_DIR, "alert.png")
        try:
            icon_code_used = tk.PhotoImage(file=icono_path)
            canvas_code_used.create_image(30, 17, anchor="nw", image=icon_code_used)
            canvas_code_used.image = icon_code_used
        except Exception as e:
            raise FileNotFoundError(f"El archivo del icono no se encontró en la ruta: {icono_path}. Error: {e}")
        
        utils.create_rounded_rectangle(canvas_code_used, 0, 66, 300, 110, radius=0, fill="#EEEEE4", outline="#EEEEE4")
        
        canvas_code_used.create_text(80, 19, text="Este código ya ha sido registrado", anchor="nw", font=("Arial", 10), fill="Black")
        canvas_code_used.create_text(115, 33, text="Por favor, utilice otro", anchor="nw", font=("Arial", 10), fill="Black")
        
        btn_code_ok = tk.Button(code_used, text="Aceptar", width=9, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=code_used.destroy)
        btn_code_ok.place(x=115, y=73)
        
        utils.aplicar_hover_a_botones([btn_code_ok])
        
    # CLIENTE REPETIDO
    def cliente_repetido(self):
        cli_rep=tk.Toplevel(self.parent)
        cli_rep.title("")
        cli_rep.geometry("300x110")
        cli_rep.resizable(False, False)
        cli_rep.configure(bg="#FFFFFF")
        cli_rep.grab_set()
        utils.centrar_ventana(cli_rep)
        cli_rep.protocol("WM_DELETE_WINDOW", lambda: None)
        
        canvas_cli_rep = tk.Canvas(cli_rep, width=300, height=110, bg="#FFFFFF", highlightthickness=0)
        canvas_cli_rep.pack()
        
        icono_path = os.path.join(ICON_DIR, "alert.png")
        try:
            icon_cli_rep = tk.PhotoImage(file=icono_path)
            canvas_cli_rep.create_image(30, 17, anchor="nw", image=icon_cli_rep)
            canvas_cli_rep.image = icon_cli_rep
        except Exception as e:
            raise FileNotFoundError(f"El archivo del icono no se encontró en la ruta: {icono_path}. Error: {e}")
        
        utils.create_rounded_rectangle(canvas_cli_rep, 0, 66, 300, 110, radius=0, fill="#EEEEE4", outline="#EEEEE4")
        
        canvas_cli_rep.create_text(100, 19, text="Este cliente ya existe", anchor="nw", font=("Arial", 10), fill="Black")
        canvas_cli_rep.create_text(95, 33, text="Intente con otro nombre", anchor="nw", font=("Arial", 10), fill="Black")
        
        btn_cli_ok = tk.Button(cli_rep, text="Aceptar", width=9, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=cli_rep.destroy)
        btn_cli_ok.place(x=115, y=73)
        
        utils.aplicar_hover_a_botones([btn_cli_ok])
        
    # COTIZACION CON CODIGO IGUAL
    def cotizacion_codigo_igual(self):
        cot_cdg_rep=tk.Toplevel(self.parent)
        cot_cdg_rep.title("")
        cot_cdg_rep.geometry("300x110")
        cot_cdg_rep.resizable(False, False)
        cot_cdg_rep.configure(bg="#FFFFFF")
        cot_cdg_rep.grab_set()
        utils.centrar_ventana(cot_cdg_rep)
        cot_cdg_rep.protocol("WM_DELETE_WINDOW", lambda: None)
        
        canvas_cdg_rep = tk.Canvas(cot_cdg_rep, width=300, height=110, bg="#FFFFFF", highlightthickness=0)
        canvas_cdg_rep.pack()
        
        icono_path = os.path.join(ICON_DIR, "alert.png")
        try:
            icon_cdg_rep = tk.PhotoImage(file=icono_path)
            canvas_cdg_rep.create_image(30, 17, anchor="nw", image=icon_cdg_rep)
            canvas_cdg_rep.image = icon_cdg_rep
        except Exception as e:
            raise FileNotFoundError(f"El archivo del icono no se encontró en la ruta: {icono_path}. Error: {e}")
        
        utils.create_rounded_rectangle(canvas_cdg_rep, 0, 66, 300, 110, radius=0, fill="#EEEEE4", outline="#EEEEE4")
        
        canvas_cdg_rep.create_text(80, 19, text="Cotización ya generada con", anchor="nw", font=("Arial", 10), fill="Black")
        canvas_cdg_rep.create_text(120, 33, text="este código", anchor="nw", font=("Arial", 10), fill="Black")
        
        btn_cdgr_ok = tk.Button(cot_cdg_rep, text="Aceptar", width=9, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=cot_cdg_rep.destroy)
        btn_cdgr_ok.place(x=115, y=73)
        
        utils.aplicar_hover_a_botones([btn_cdgr_ok])
        
    # VALIDAR EL COSTO DE LOS MATERIALES
    def validar_numeros(self):
        val_prec=tk.Toplevel(self.parent)
        val_prec.title("")
        val_prec.geometry("300x110")
        val_prec.resizable(False, False)
        val_prec.configure(bg="#FFFFFF")
        val_prec.grab_set()
        utils.centrar_ventana(val_prec)
        val_prec.protocol("WM_DELETE_WINDOW", lambda: None)
        
        canvas_val_prec = tk.Canvas(val_prec, width=300, height=110, bg="#FFFFFF", highlightthickness=0)
        canvas_val_prec.pack()
        
        icono_path = os.path.join(ICON_DIR, "alert.png")
        try:
            icon_val_prec = tk.PhotoImage(file=icono_path)
            canvas_val_prec.create_image(30, 17, anchor="nw", image=icon_val_prec)
            canvas_val_prec.image = icon_val_prec
        except Exception as e:
            raise FileNotFoundError(f"El archivo del icono no se encontró en la ruta: {icono_path}. Error: {e}")
        
        utils.create_rounded_rectangle(canvas_val_prec, 0, 66, 300, 110, radius=0, fill="#EEEEE4", outline="#EEEEE4")
        
        canvas_val_prec.create_text(80, 19, text="Por favor, ingrese un número", anchor="nw", font=("Arial", 10), fill="Black")
        canvas_val_prec.create_text(135, 33, text="válido", anchor="nw", font=("Arial", 10), fill="Black")
        
        btn_val_ok = tk.Button(val_prec, text="Aceptar", width=9, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=val_prec.destroy)
        btn_val_ok.place(x=115, y=73)
        
        utils.aplicar_hover_a_botones([btn_val_ok])


class generator_cot:
    def __init__(self, root, gen_cot):
        
        self.root = root
        self.gen_cot = gen_cot
        self.root.withdraw()
        
        self.gen_cot = gen_cot
        self.gen_cot.title("Generar Cotización")
        self.gen_cot.geometry("840x714")
        self.gen_cot.resizable(False, False)
        self.gen_cot.configure(bg="#373737")
        utils.centrar_ventana(self.gen_cot)
        self.alerta = alertas(gen_cot)
        self.gen_cot.protocol("WM_DELETE_WINDOW", lambda: None)
        
        self.TEMPLATE_FILE = os.path.join(COT_PLAN, "plantilla_cotizacion.docx")
        if not os.path.exists(self.TEMPLATE_FILE):
            raise FileNotFoundError(f"No se encontró la plantilla: {self.TEMPLATE_FILE}")

        self.subtotal_general = 0.0
        self.selected_item = None
        
        # Definir las cuentas
        self.cuentas_soles = {
            'Nro de Cuenta': '191-2539938-1-47',
            'Cuenta Interbancaria': 'Cuenta Soles IBAN'
        }
        self.cuentas_dolares = {
            'Nro de Cuenta': '12345678910987654321',
            'Cuenta Interbancaria': 'Cuenta Dólares IBAN'
        }
        
        canvas_coti = tk.Canvas(gen_cot, width=840, height=714, bg="#373737", highlightthickness=0)
        canvas_coti.pack()
        
        utils.create_rounded_rectangle(canvas_coti, 10, 10, 830, 252, radius=10, fill="#959595", outline="#959595")
        utils.create_rounded_rectangle(canvas_coti, 10, 262, 830, 388, radius=10, fill="#959595", outline="#959595")
        #utils.create_rounded_rectangle(canvas_coti, 10, 438, 830, 586, radius=10, fill="#959595", outline="#959595")
        utils.create_rounded_rectangle(canvas_coti, 10, 596, 830, 664, radius=10, fill="#959595", outline="#959595")
        
        
        canvas_coti.create_text(20, 20, text="Nro de Cotización", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_coti, 20, 38, 170, 68, radius=10, fill="white", outline="#959595")
        self.input_gcoti = tk.Entry(gen_cot, font=("Arial", 11), bd=0)
        self.input_gcoti.place(x=25, y=43, width=140, height=20)
        num_cotizacion = self.generar_numero_cotizacion()
        self.input_gcoti.insert(0, num_cotizacion)
        
        canvas_coti.create_text(180, 20, text="Fecha", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_coti, 180, 38, 410, 68, radius=10, fill="white", outline="#959595")
        self.input_gfecha = tk.Entry(gen_cot, font=("Arial", 11), bd=0)
        self.input_gfecha.place(x=185, y=43, width=220, height=20)
        
        canvas_coti.create_text(420, 20, text="Cliente / Empresa", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_coti, 420, 38, 620, 68, radius=10, fill="white", outline="#959595")
        self.input_gpersona = tk.Entry(gen_cot, font=("Arial", 11), bd=0)
        self.input_gpersona.place(x=425, y=43, width=190, height=20)
        
        self.lista_sugerencias = tk.Listbox(self.gen_cot)
        self.lista_sugerencias.place_forget()
        self.lista_sugerencias.bind("<<ListboxSelect>>", self.on_select_cliente)
        
        self.input_gpersona.bind("<KeyRelease>", self.update_suggestions)
        
        canvas_coti.create_text(630, 20, text="Área de Trabajo", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        self.cbo_area = ttk.Combobox(gen_cot, values=["Escoja una Opción", "Ejemplo1"], state="readonly", font=("Raleway", 10))
        self.cbo_area.place(x=630, y=38, width=190, height=31)
        self.cbo_area.current(0)
        
        
        canvas_coti.create_text(20, 78, text="Persona de Contacto", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        self.cbo_persona = ttk.Combobox(gen_cot, values=["Escoja una Opción", "Ejemplo1"], state="readonly", font=("Raleway", 10))
        self.cbo_persona.place(x=20, y=96, width=230, height=31)
        self.cbo_persona.current(0)
        
        canvas_coti.create_text(260, 78, text="Título del Servicio", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_coti, 260, 96, 820, 126, radius=10, fill="white", outline="#959595")
        self.input_gtitulo = tk.Entry(gen_cot, font=("Arial", 11), bd=0)
        self.input_gtitulo.place(x=265, y=101, width=550, height=20)
        
        
        canvas_coti.create_text(20, 136, text="Tiempo de Ejecución", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_coti, 20, 154, 400, 184, radius=10, fill="white", outline="#959595")
        self.input_gtiempo = tk.Entry(gen_cot, font=("Arial", 11), bd=0)
        self.input_gtiempo.place(x=25, y=159, width=370, height=20)
        
        canvas_coti.create_text(410, 136, text="Forma de Pago", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_coti, 410, 154, 630, 184, radius=10, fill="white", outline="#959595")
        self.input_gpago = tk.Entry(gen_cot, font=("Arial", 11), bd=0)
        self.input_gpago.place(x=415, y=159, width=210, height=20)
        
        canvas_coti.create_text(640, 136, text="IGV", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        self.cbo_igv = ttk.Combobox(gen_cot, values=["SI", "NO"], state="readonly", font=("Raleway", 10))
        self.cbo_igv.place(x=640, y=154, width=50, height=31)
        self.cbo_igv.current(0)
        self.cbo_igv.bind("<<ComboboxSelected>>", self.actualizar_totales)
        
        canvas_coti.create_text(700, 136, text="Tipo de Moneda", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        self.cbo_moneda = ttk.Combobox(gen_cot, values=["PEN", "USD"], state="readonly", font=("Raleway", 10))
        self.cbo_moneda.place(x=700, y=154, width=120, height=31)
        self.cbo_moneda.current(0)
        
        canvas_coti.create_text(20, 194, text="Tipo de Cuenta", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        self.cbo_cuenta = ttk.Combobox(gen_cot, values=["Cuenta Soles", "Cuenta Dólares"], state="readonly", font=("Raleway", 10))
        self.cbo_cuenta.place(x=20, y=212, width=180, height=31)
        self.cbo_cuenta.current(0)
        self.cbo_cuenta.bind("<<ComboboxSelected>>", self.actualizar_cuentas)
        
        canvas_coti.create_text(210, 194, text="Nro de Cuenta", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_coti, 210, 212, 510, 242, radius=10, fill="white", outline="#959595")
        self.input_gcuenta = tk.Entry(gen_cot, font=("Arial", 11), bd=0)
        self.input_gcuenta.place(x=215, y=217, width=290, height=20)
        
        canvas_coti.create_text(520, 194, text="Cuenta Interbancaria", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_coti, 520, 212, 820, 242, radius=10, fill="white", outline="#959595")
        self.input_gcuentabanc = tk.Entry(gen_cot, font=("Arial", 11), bd=0)
        self.input_gcuentabanc.place(x=525, y=217, width=290, height=20)
        
        
        canvas_coti.create_text(20, 272, text="Descripción", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_coti, 20, 290, 820, 320, radius=10, fill="white", outline="#959595")
        self.input_gdesc = tk.Entry(gen_cot, font=("Arial", 11), bd=0)
        self.input_gdesc.place(x=25, y=295, width=790, height=20)
        
        canvas_coti.create_text(20, 330, text="Material(es)", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_coti, 20, 348, 485, 378, radius=10, fill="white", outline="#959595")
        self.input_gmaterial = tk.Entry(gen_cot, font=("Arial", 11), bd=0)
        self.input_gmaterial.place(x=25, y=353, width=455, height=20)
        
        canvas_coti.create_text(495, 330, text="Unidad(es)", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_coti, 495, 348, 575, 378, radius=10, fill="white", outline="#959595")
        self.input_gunidad = tk.Entry(gen_cot, font=("Arial", 11), bd=0)
        self.input_gunidad.place(x=500, y=353, width=70, height=20)
        self.cbo_unidad = ttk.Combobox(gen_cot, values=["JUEGO", "PIEZA"], state="readonly", font=("Raleway", 10))
        self.cbo_unidad.place(x=585, y=348, width=100, height=31)
        self.cbo_unidad.current(0)
        
        canvas_coti.create_text(695, 330, text="Precio Unit.", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_coti, 695, 348, 820, 378, radius=10, fill="white", outline="#959595")
        self.input_gpreciou = tk.Entry(gen_cot, font=("Arial", 11), bd=0)
        self.input_gpreciou.place(x=700, y=353, width=115, height=20)
        
        
        canvas_coti.create_text(20, 606, text="Dscto. (%)", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_coti, 20, 624, 130, 654, radius=10, fill="white", outline="#959595")
        self.input_gdescuento = tk.Entry(gen_cot, font=("Arial", 11), bd=0)
        self.input_gdescuento.place(x=25, y=629, width=100, height=20)
        self.descuento_var = tk.StringVar()
        self.input_gdescuento.config(textvariable=self.descuento_var)
        self.descuento_var.trace('w', self.actualizar_totales)
        
        canvas_coti.create_text(350, 606, text="SubTotal", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_coti, 350, 624, 460, 654, radius=10, fill="white", outline="#959595")
        self.input_gsubtot = tk.Entry(gen_cot, font=("Arial", 11), bd=0)
        self.input_gsubtot.place(x=355, y=629, width=100, height=20)
        
        canvas_coti.create_text(470, 606, text="IGV", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_coti, 470, 624, 580, 654, radius=10, fill="white", outline="#959595")
        self.input_gigv = tk.Entry(gen_cot, font=("Arial", 11), bd=0)
        self.input_gigv.place(x=475, y=629, width=100, height=20)
        
        canvas_coti.create_text(590, 606, text="Dscto.", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_coti, 590, 624, 700, 654, radius=10, fill="white", outline="#959595")
        self.input_gdescto = tk.Entry(gen_cot, font=("Arial", 11), bd=0)
        self.input_gdescto.place(x=595, y=629, width=100, height=20)
        
        canvas_coti.create_text(710, 606, text="Total", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_coti, 710, 624, 820, 654, radius=10, fill="white", outline="#959595")
        self.input_gtotal = tk.Entry(gen_cot, font=("Arial", 11), bd=0)
        self.input_gtotal.place(x=715, y=629, width=100, height=20)
        
        
        btn_ag = tk.Button(gen_cot, text="Agregar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=self.agregar_material)
        btn_ag.place(x=10, y=398)
        
        btn_ed = tk.Button(gen_cot, text="Editar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=self.editar_material)
        btn_ed.place(x=120, y=398)
        
        btn_del = tk.Button(gen_cot, text="Eliminar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=self.eliminar_material)
        btn_del.place(x=230, y=398)
            
        btn_canc = tk.Button(gen_cot, text="Cancelar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=self.cancelar)
        btn_canc.place(x=10, y=674)

        btn_gen = tk.Button(gen_cot, text="Generar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=self.guardar_registro)
        btn_gen.place(x=120, y=674)
        
        utils.aplicar_hover_a_botones([btn_ag, btn_ed, btn_del, btn_canc, btn_gen])
        
        self.t_material = ttk.Treeview(gen_cot, columns=("desc", "mat", "und", "p_unit", "s_total"), show="headings", style="Custom.Treeview")
        self.t_material.place(x=10, y=438, width=821, height=149)
                                                    # - 155
        
        self.t_material.heading("desc", text="Descripción")
        self.t_material.heading("mat", text="Material")
        self.t_material.heading("und", text="Unidad(es)")
        self.t_material.heading("p_unit", text="Precio Unit.")
        self.t_material.heading("s_total", text="SubTotal")
        
        self.t_material.column("desc", anchor="center", width=316, stretch=False)
        self.t_material.column("mat", anchor="center", width=192, stretch=False)
        self.t_material.column("und", anchor="center", width=99, stretch=False)
        self.t_material.column("p_unit", anchor="center", width=99, stretch=False)
        self.t_material.column("s_total", anchor="center", width=99, stretch=False)
        
        self.t_material.bind("<Double-1>", self.editar_material)
        
        scrollbar_vertical = ttk.Scrollbar(gen_cot, orient="vertical", command=self.t_material.yview)
        self.t_material.configure(yscrollcommand=scrollbar_vertical.set)
        scrollbar_vertical.place(x=817, y=438, height=149)
        
    def generar_numero_cotizacion(self):
        archivo_numero = 'numero_cotizacion.txt'
        try:
            with open(archivo_numero, 'r') as f:
                ultimo_numero = int(f.read())
        except FileNotFoundError:
            ultimo_numero = 1500  # Número inicial si no existe el archivo
        nuevo_numero = ultimo_numero + 2
        with open(archivo_numero, 'w') as f:
            f.write(str(nuevo_numero))
        # Formato del número de cotización
        numero_cotizacion_formateado = f"002/{nuevo_numero}"
        return numero_cotizacion_formateado
    
    def actualizar_cuentas(self, event=None):
        tipo_cuenta = self.cbo_cuenta.get()
        if tipo_cuenta == 'Cuenta Soles':
            cuentas = self.cuentas_soles
        elif tipo_cuenta == 'Cuenta Dólares':
            cuentas = self.cuentas_dolares
        else:
            cuentas = {'Nro de Cuenta': '', 'Cuenta Interbancaria': ''}

        self.input_gcuenta.delete(0, tk.END)
        self.input_gcuenta.insert(0, cuentas['Nro de Cuenta'])

        self.input_gcuentabanc.delete(0, tk.END)
        self.input_gcuentabanc.insert(0, cuentas['Cuenta Interbancaria'])
        
    def agregar_material(self):
        descripcion = self.input_gdesc.get().strip()
        material = self.input_gmaterial.get().strip()
        unidades = self.input_gunidad.get().strip()
        tipo_unidad = self.cbo_unidad.get().strip()
        precio_unitario = self.input_gpreciou.get().strip()
        
        # Validar entradas
        if not descripcion or not material or not unidades or not precio_unitario or not tipo_unidad:
            self.alerta.question_datos()
            return
        try:
            unidades = int(unidades)
            precio_unitario = float(precio_unitario)
        except ValueError:
            self.alerta.validar_numeros()
            return

        subtotal = unidades * precio_unitario
        
        # Insertar en la tabla
        self.t_material.insert("", "end", values=(
            descripcion,
            material,
            f"{unidades} {tipo_unidad}",
            f"{precio_unitario:.2f}",
            f"{subtotal:.2f}"
        ))
        
        # Recalcular el subtotal general
        self.recalcular_subtotal_general()

        # Actualizar los totales
        self.actualizar_totales()

        # Limpiar campos de entrada
        self.limpiar_campos_material()

        # Mostrar alerta de registro exitoso
        self.alerta.registro_confirm()
        
    def guardar_registro(self):
        incluir_igv = self.cbo_igv.get() == "SI"
        moneda = self.cbo_moneda.get()

        # Validar que el subtotal no esté vacío
        subtotal_str = self.input_gsubtot.get()
        if not subtotal_str:
            self.alerta.question_datos()
            return
        
        # Obtener valores de descuento y totales
        try:
            descuento_monto_str = self.input_gdescto.get().strip()
            if descuento_monto_str == '':
                descuento_monto = 0.0
            else:
                descuento_monto = float(descuento_monto_str)

            descuento_porc_str = self.descuento_var.get().strip()
            if descuento_porc_str == '':
                descuento_porc = 0.0
            else:
                descuento_porc = float(descuento_porc_str)

            subtotal_descuento = float(self.input_gsubtot.get())
            igv = float(self.input_gigv.get())
            total = float(self.input_gtotal.get())
        except ValueError:
            self.alerta.validar_numeros()
            return
        
        # Recopilar datos de la cotización
        data_cotizacion = {
            'N° de Cotización': self.input_gcoti.get(),
            'Fecha': self.input_gfecha.get(),
            'Persona de contacto': self.cbo_persona.get(),
            'Empresa': self.input_gpersona.get(),
            'Área de Trabajo': self.cbo_area.get(),
            'Servicio': self.input_gtitulo.get(),
            'Subtotal': self.subtotal_general,  # Subtotal antes del descuento
            'Descuento (%)': descuento_porc,
            'Descuento': descuento_monto,
            'Subtotal con Descuento': subtotal_descuento,
            'IGV (18%)': igv,
            'Total': total,
            'Tiempo de Ejecución': self.input_gtiempo.get(),
            'Forma de Pago': self.input_gpago.get(),
            'Cuenta': self.input_gcuenta.get(),
            'Cuenta Interbancaria': self.input_gcuentabanc.get(),
            'Tipo de Cuenta': self.cbo_cuenta.get(),
            'Tipo de Pago': 'Cuenta Corriente BCP soles' if self.cbo_cuenta.get() == 'Cuenta Soles' else 'Cuenta Corriente BCP dólares',
            'Estado': 'Pendiente',
            'confirmacion': 'Pendiente'
        }
        
        # Obtener los materiales de la tabla
        materiales = []
        for item in self.t_material.get_children():
            valores = self.t_material.item(item, "values")
            unidades_str, tipo_unidad = valores[2].split()
            materiales.append({
                "Descripción": valores[0],
                "Material": valores[1],
                "Unidades": int(unidades_str),
                "Tipo Unidad": tipo_unidad,
                "Precio Unitario": float(valores[3]),
                "Subtotal Unidades": float(valores[4])
            })
            
        # Generar el documento Word
        output_file = self.crear_documento_word(data_cotizacion, materiales, moneda, incluir_igv)

        # Mostrar alerta de confirmación
        self.alerta.confirmacion_cotizacion()

        # Cerrar la ventana actual y volver a la principal
        self.gen_cot.destroy()
        self.root.deiconify()
        
        # Confirmar la generación de la cotización
        def confirmar_generacion():
            # Generar el documento Word
            output_file = self.crear_documento_word(data_cotizacion, materiales, moneda, incluir_igv)

            if output_file:
                # Mostrar alerta de registro exitoso
                self.alerta.confirmacion_cotizacion()

                # Cerrar la ventana actual y volver a la principal
                self.gen_cot.destroy()
                self.root.deiconify()
            else:
                # Mostrar alerta si no se pudo generar el documento
                self.alerta.mostrar_mensaje("No se pudo generar la cotización.")

        # Mostrar alerta de confirmación
        self.alerta.generator_cotizacion(confirmar_generacion)
        
    def eliminar_material(self):
        selected_items = self.t_material.selection()
        if not selected_items:
            self.alerta.mostrar_mensaje("Por favor, seleccione una fila para eliminar.")
            return

        # Mostrar la alerta de confirmación sin callback
        self.alerta.quest_mat()

        # Eliminar las filas seleccionadas
        for item in selected_items:
            self.t_material.delete(item)
        
        # Recalcular el subtotal general y actualizar totales
        self.recalcular_subtotal_general()
        self.actualizar_totales()
        
        # Mostrar alerta de eliminación exitosa
        self.alerta.material_delete()
        
    def cancelar(self):
        self.gen_cot.destroy()
        self.root.deiconify()
    
    
    def editar_material(self):
        
        selected_items = self.t_material.selection()
        if not selected_items:
            self.alerta.seleccionar_fila()
            return

        item_id = selected_items[0]
        item_values = self.t_material.item(item_id, "values")
        
        ed_material = tk.Toplevel(self.gen_cot)
        ed_material.title("Editar Material")
        ed_material.geometry("600x244")
        ed_material.resizable(False, False)
        ed_material.configure(bg="#373737")
        ed_material.grab_set()
        utils.centrar_ventana(ed_material)
        self.alerta = alertas(ed_material)
        ed_material.protocol("WM_DELETE_WINDOW", lambda: None)
        
        canvas_mat = tk.Canvas(ed_material, width=600, height=244, bg="#373737", highlightthickness=0)
        canvas_mat.pack()
        
        utils.create_rounded_rectangle(canvas_mat, 10, 10, 590, 194, radius=10, fill="#959595", outline="#959595")
        
        canvas_mat.create_text(20, 20, text="Descripción", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_mat, 20, 38, 580, 68, radius=10, fill="white", outline="#959595")
        inpt_ed_desc = tk.Entry(ed_material, font=("Arial", 11), bd=0)
        inpt_ed_desc.place(x=25, y=43, width=550, height=20)
        
        canvas_mat.create_text(20, 78, text="Material", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_mat, 20, 96, 580, 126, radius=10, fill="white", outline="#959595")
        inpt_ed_mat = tk.Entry(ed_material, font=("Arial", 11), bd=0)
        inpt_ed_mat.place(x=25, y=101, width=550, height=20)
    
        canvas_mat.create_text(20, 136, text="Unidad(es)", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_mat, 20, 154, 103, 184, radius=10, fill="white", outline="#959595")
        inpt_ed_unidad = tk.Entry(ed_material, font=("Arial", 11), bd=0)
        inpt_ed_unidad.place(x=25, y=159, width=73, height=20)
        cbo_ed_unidad = ttk.Combobox(ed_material, values=["JUEGO", "PIEZA"], state="readonly", font=("Raleway", 10))
        cbo_ed_unidad.place(x=113, y=154, width=90, height=31)
        cbo_ed_unidad.current(0)
        
        canvas_mat.create_text(213, 136, text="Precio Unit.", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_mat, 213, 154, 343, 184, radius=10, fill="white", outline="#959595")
        inpt_ed_precio = tk.Entry(ed_material, font=("Arial", 11), bd=0)
        inpt_ed_precio.place(x=218, y=159, width=120, height=20)
        
        canvas_mat.create_text(353, 136, text="Sub Total", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_mat, 353, 154, 483, 184, radius=10, fill="white", outline="#959595")
        inpt_ed_total = tk.Entry(ed_material, font=("Arial", 11), bd=0)
        inpt_ed_total.place(x=358, y=159, width=120, height=20)
        
        btn_ed_canc = tk.Button(ed_material, text="Cancelar", width=15, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=ed_material.destroy)
        btn_ed_canc.place(x=10, y=204)
        
        def guardar_edicion():
            descripcion = inpt_ed_desc.get().strip()
            material = inpt_ed_mat.get().strip()
            try:
                unidades = int(inpt_ed_unidad.get().strip())
                precio_unitario = float(inpt_ed_precio.get().strip())
            except ValueError:
                self.alerta.mostrar_mensaje("Por favor, ingrese un número válido en 'Unidad(es)' y 'Precio Unit.'")
                return
            tipo_unidad = cbo_ed_unidad.get().strip()
            subtotal = unidades * precio_unitario
            
            # Actualizar el Treeview
            self.t_material.item(item_id, values=(
                descripcion,
                material,
                f"{unidades} {tipo_unidad}",
                f"{precio_unitario:.2f}",
                f"{subtotal:.2f}"
            ))
            
            # Recalcular el subtotal general y actualizar totales
            self.recalcular_subtotal_general()
            self.actualizar_totales()
            
            # Cerrar la ventana de edición
            ed_material.destroy()
            
            # Mostrar alerta de edición exitosa
            self.alerta.registro_confirm()

        btn_ed_save = tk.Button(ed_material, text="Guardar", width=15, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=guardar_edicion)
        btn_ed_save.place(x=134, y=204)
        
        utils.aplicar_hover_a_botones([btn_ed_canc, btn_ed_save])
        
    def actualizar_totales(self, *args):
        incluir_igv = self.cbo_igv.get() == "SI"

        # Obtener el porcentaje de descuento
        descuento_porc_str = self.descuento_var.get().strip()
        if descuento_porc_str == '':
            descuento_porc = 0.0
        else:
            try:
                descuento_porc = float(descuento_porc_str)
            except ValueError:
                self.alerta.mostrar_mensaje("Por favor, ingrese un porcentaje válido en 'Dscto. (%)'")
                descuento_porc = 0.0  # Si la entrada no es válida, se asume 0%

        descuento_monto, subtotal_descuento, igv, total = self.calcular_totales(
            self.subtotal_general, descuento_porc, incluir_igv
        )
        
        # Actualizar los campos en la interfaz
        self.input_gsubtot.config(state='normal')
        self.input_gsubtot.delete(0, tk.END)
        self.input_gsubtot.insert(0, f"{subtotal_descuento:.2f}")
        self.input_gsubtot.config(state='readonly')

        self.input_gigv.config(state='normal')
        self.input_gigv.delete(0, tk.END)
        self.input_gigv.insert(0, f"{igv:.2f}")
        self.input_gigv.config(state='readonly')

        self.input_gtotal.config(state='normal')
        self.input_gtotal.delete(0, tk.END)
        self.input_gtotal.insert(0, f"{total:.2f}")
        self.input_gtotal.config(state='readonly')

        self.input_gdescto.config(state='normal')
        self.input_gdescto.delete(0, tk.END)
        self.input_gdescto.insert(0, f"{descuento_monto:.2f}")
        self.input_gdescto.config(state='readonly')
        
    def calcular_totales(self, subtotal_general, descuento_porc, incluir_igv):
        descuento_monto = subtotal_general * descuento_porc / 100.0
        subtotal_descuento = subtotal_general - descuento_monto
        igv = subtotal_descuento * 0.18 if incluir_igv else 0.0
        total = subtotal_descuento + igv
        return descuento_monto, subtotal_descuento, igv, total

    def recalcular_subtotal_general(self):
        self.subtotal_general = 0.0
        for item in self.t_material.get_children():
            item_values = self.t_material.item(item, "values")
            item_subtotal = float(item_values[4])
            self.subtotal_general += item_subtotal
            
    def limpiar_campos_material(self):
        self.input_gdesc.delete(0, tk.END)
        self.input_gmaterial.delete(0, tk.END)
        self.input_gunidad.delete(0, tk.END)
        self.input_gpreciou.delete(0, tk.END)
        self.cbo_unidad.set("JUEGO")  # Restablecer al valor predeterminado

    def crear_documento_word(self, data, materiales, moneda, incluir_igv):
        # Definir el símbolo de moneda
        simbolo_moneda = "S/" if moneda == "PEN" else "$"

        # Generar el subtotal en letras
        subtotal_para_letras = data['Subtotal']
        subtotal_en_letras = self.convertir_numero_a_texto(subtotal_para_letras, moneda, incluir_igv)
        
        # Obtener el tipo de pago basado en el tipo de cuenta
        if data['Tipo de Cuenta'] == 'Cuenta Soles':
            tipo_pago = "Cuenta Corriente BCP Soles: 191-2539938-1-47"
        elif data['Tipo de Cuenta'] == 'Cuenta Dólares':
            tipo_pago = "Cuenta Corriente BCP Dólares: 12345678910987654321"
        else:
            tipo_pago = ''
            
        # Cargar la plantilla de Word
        try:
            doc = Document(self.TEMPLATE_FILE)  # Cargar plantilla
        except Exception as e:
            self.alerta.mostrar_mensaje(f"No se pudo cargar la plantilla: {e}")
            return
        
        # Reemplazar los marcadores en el documento
        for paragraph in doc.paragraphs:
            if '[N° DE COTIZACIÓN]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[N° DE COTIZACIÓN]', data['N° de Cotización'])
            if '[FECHA]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[FECHA]', data['Fecha'])
            if '[NOMBRE_CLIENTE]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[NOMBRE_CLIENTE]', data['Persona de contacto'])
            if '[EMPRESA]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[EMPRESA]', data['Empresa'])
            if '[SERVICIO]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[SERVICIO]', data['Servicio'])
            if '[TOTAL_EN_LETRAS]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[TOTAL_EN_LETRAS]', subtotal_en_letras)
            if '[TIEMPO_EJECUCION]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[TIEMPO_EJECUCION]', data['Tiempo de Ejecución'])
            if '[FORMA_PAGO]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[FORMA_PAGO]', data['Forma de Pago'])
            if '[NOTA_IGV]' in paragraph.text:
                nota_igv = "Los precios incluyen IGV." if incluir_igv else "Los precios no incluyen IGV."
                paragraph.text = paragraph.text.replace('[NOTA_IGV]', nota_igv)
            if '[TIPO_PAGO]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[TIPO_PAGO]', tipo_pago)

        # Buscar la tabla en el documento donde se insertarán los servicios/materiales
        table = doc.tables[0]
        
        # Añadir servicios/materiales a la tabla
        for i, servicio in enumerate(materiales):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i + 1)  # Número del servicio/material
            row_cells[1].text = servicio['Descripción']  # Descripción
            row_cells[2].text = servicio['Material']  # Material
            row_cells[3].text = f"{servicio['Unidades']} {servicio['Tipo Unidad']}"  # Unidades
            row_cells[4].text = f"{simbolo_moneda} {servicio['Precio Unitario']:.2f}"  # Precio Unitario
            row_cells[5].text = f"{simbolo_moneda} {servicio['Subtotal Unidades']:.2f}"  # Subtotal de las unidades

        # Añadir el subtotal antes del descuento
        row = table.add_row().cells
        row[4].text = "Subtotal:"
        row[5].text = f"{simbolo_moneda} {data['Subtotal']:.2f}"
        
        # Añadir el descuento si existe
        if data['Descuento'] > 0:
            row = table.add_row().cells
            row[4].text = f"Descuento:"
            row[5].text = f"-{simbolo_moneda} {data['Descuento']:.2f}"

        # Añadir IGV y total
        if incluir_igv:
            row = table.add_row().cells
            row[4].text = "IGV (18%):"
            row[5].text = f"{simbolo_moneda} {data['IGV (18%)']:.2f}"

        # Total
        row = table.add_row().cells
        row[4].text = "Total:"
        row[5].text = f"{simbolo_moneda} {data['Total']:.2f}"
        
        # Crear la carpeta si no existe para guardar las cotizaciones
        carpeta_cotizaciones = os.path.join('clientes', data['Empresa'].replace('/', '_').replace('\\', '_'), 'cotizaciones')
        os.makedirs(carpeta_cotizaciones, exist_ok=True)

        # Guardar el documento en una ubicación con el número de cotización
        cotizacion_num = data['N° de Cotización'].replace('/', '_').replace('\\', '_')
        output_file = os.path.join(carpeta_cotizaciones, f"Cotizacion_{cotizacion_num}.docx")
        doc.save(output_file)
        
        # Retornar la ruta del archivo
        return output_file
    
    def convertir_numero_a_texto(self, numero, moneda, incluir_igv):
        parte_entera = int(numero)
        parte_decimal = round((numero - parte_entera) * 100)

        parte_entera_texto = num2words(parte_entera, lang='es', to='cardinal').capitalize()
        parte_decimal_texto = num2words(parte_decimal, lang='es', to='cardinal')

        if moneda == "PEN":
            moneda_texto = "soles"
            simbolo_moneda = "S/"
        else:
            moneda_texto = "dólares americanos"
            simbolo_moneda = "$"

        # Ajustar el texto según si incluye IGV o no
        igv_texto = " + IGV" if incluir_igv else ""

        texto = f"{simbolo_moneda} {numero:.2f} ({parte_entera_texto} con {parte_decimal_texto} {moneda_texto}){igv_texto}"
        return texto
    
    def on_select_cliente(self, event):
        try:
            seleccion = self.lista_sugerencias.curselection()
            if seleccion:
                indice = seleccion[0]
                valor = self.lista_sugerencias.get(indice)
                self.input_gpersona.delete(0, tk.END)
                self.input_gpersona.insert(0, valor)
                self.lista_sugerencias.place_forget()
                # Aquí puedes agregar lógica adicional para cargar datos del cliente seleccionado si es necesario
        except Exception as e:
            self.alerta.mostrar_mensaje(f"Error al seleccionar el cliente: {e}")

    def update_suggestions(self, event):
        texto = self.input_gpersona.get().strip()
        if texto == '':
            self.lista_sugerencias.place_forget()
            return
        
        # Aquí deberías implementar la lógica para buscar clientes que coincidan con 'texto'.
        # Por ejemplo, podrías tener una lista de clientes predefinida o cargarla desde un archivo.
        # Para este ejemplo, usaremos una lista estática.
        clientes = ['Cliente A', 'Cliente B', 'Cliente C', 'Cliente D']
        resultados = [cliente for cliente in clientes if texto.lower() in cliente.lower()]
        
        if resultados:
            self.lista_sugerencias.delete(0, tk.END)
            for r in resultados:
                self.lista_sugerencias.insert(tk.END, r)
            
            # Posicionar el Listbox debajo del Entry
            x = self.input_gpersona.winfo_x()
            y = self.input_gpersona.winfo_y() + self.input_gpersona.winfo_height()
            self.lista_sugerencias.place(x=x, y=y, width=self.input_gpersona.winfo_width())
        else:
            self.lista_sugerencias.place_forget()        


class clientes:
    def __init__(self, root, vent_clientes):
        
        self.root = root
        self.vent_clientes = vent_clientes
        self.root.withdraw()
        
        self.vent_clientes = vent_clientes
        self.vent_clientes.title("Registro de Clientes")
        self.vent_clientes.geometry("1200x698")
        self.vent_clientes.resizable(False, False)
        self.vent_clientes.configure(bg="#373737")
        utils.centrar_ventana(self.vent_clientes)
        
        self.alerta = alertas(vent_clientes)

        self.vent_clientes.protocol("WM_DELETE_WINDOW", lambda: None)
        
        canvas_cliente = tk.Canvas(vent_clientes, width=1200, height=698, bg="#373737", highlightthickness=0)
        canvas_cliente.pack()
        
        utils.create_rounded_rectangle(canvas_cliente, 10, 10, 300, 554, radius=10, fill="#959595", outline="#959595")
        utils.create_rounded_rectangle(canvas_cliente, 10, 564, 300, 688, radius=10, fill="#959595", outline="#959595")
        #utils.create_rounded_rectangle(canvas_cliente, 310, 80, 968, 408, radius=10, fill="#959595", outline="#959595")
        
        canvas_cliente.create_text(20, 22, text="Opciones", anchor="nw", font=("Raleway", 20, "bold"), fill="White")
        
        btn_reg_cliente = tk.Button(vent_clientes, text="Registrar Cliente", width=37, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=self.registrar_cliente)
        btn_reg_cliente.place(x=22, y=80)
        
        btn_ed_cli = tk.Button(vent_clientes, text="Editar cliente", width=37, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=self.detalle_cliente)
        btn_ed_cli.place(x=22, y=125)
        
        btn_menu = tk.Button(vent_clientes, text="Volver al inicio", width=37, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_menu.place(x=22, y=170)
        
        btn_sig_cli = tk.Button(vent_clientes, text="Siguiente", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_sig_cli.place(x=1090, y=658)
        
        btn_atras_cli = tk.Button(vent_clientes, text="Anterior", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_atras_cli.place(x=980, y=658)
        
        utils.aplicar_hover_a_botones([btn_reg_cliente, btn_ed_cli, btn_menu, btn_sig_cli, btn_atras_cli])
        
        canvas_cliente.create_text(20, 576, text="Filtros", anchor="nw", font=("Raleway", 20, "bold"), fill="White")
        
        canvas_cliente.create_text(20, 630, text="Por Fecha", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        
        cbo_date = ttk.Combobox(vent_clientes, values=["Todos los registros", "Última semana", "Últimas 2 semanas", "Último mes", "Últimos 3 meses", "Últimos 6 meses", "Último año"], state="readonly", font=("Raleway", 10))
        cbo_date.place(x=20, y=648, width=270, height=30)
        cbo_date.current(0)
        
        canvas_cliente.create_text(425, 22, text="REGISTRO DE CLIENTES", anchor="nw", font=("Raleway", 20, "bold"), fill="White")
        
        search_canvas_cliente = tk.Canvas(vent_clientes, width=350, height=40, bg="#373737", highlightthickness=0)
        search_canvas_cliente.place(x=840, y=20)
                
        utils.create_rounded_rectangle(search_canvas_cliente, 0, 0, 350, 40, radius=10, fill="white", outline="gray")
        search_canvas_cliente.create_line(315, 7, 315, 34, fill="gray", width=2)
        
        search_icon_path = os.path.join(ICON_DIR, "search.png")
        try:
            search_icon_cliente = tk.PhotoImage(file=search_icon_path)
            search_icon_id_cliente = search_canvas_cliente.create_image(321, 8, anchor="nw", image=search_icon_cliente)
            search_canvas_cliente.image = search_icon_cliente
        except Exception as e:
            raise FileNotFoundError(f"El archivo del icono no se encontró en la ruta: {search_icon_path}. Error: {e}")
        
        search_canvas_cliente.tag_bind(search_icon_id_cliente, "<Button-1>", lambda e: self.alerta.no_datos())
                
        search_entry_cliente = tk.Entry(search_canvas_cliente, font=("Arial", 13), width=40, bd=0, relief="flat", fg='grey')
        search_entry_cliente.insert(0, "Buscar...")
        search_entry_cliente.bind("<FocusIn>", lambda event: utils.clear_placeholder(event, search_entry_cliente))
        search_entry_cliente.bind("<FocusOut>", lambda event: utils.placeholder_search(event, search_entry_cliente))
        search_entry_cliente.place(x=6, y=7, width=301, height=27)
        
        cbo_page_cliente = ttk.Combobox(vent_clientes, values=["1", "2", "1000"], state="readonly", font=("Arial", 10))
        cbo_page_cliente.place(x=310, y=658, width=70, height=30)
        cbo_page_cliente.current(0)

        t_cliente = ttk.Treeview(vent_clientes, columns=("id", "razon", "ruc", "fecha",), show="headings", style="Custom.Treeview")
        t_cliente.place(x=310, y=80, width=881, height=569)
        
        t_cliente.heading("id", text="ID")
        t_cliente.heading("razon", text="Razón Social / Cliente / Empresa")
        t_cliente.heading("ruc", text="RUC / DNI")
        t_cliente.heading("fecha", text="Fecha")
        
        t_cliente.column("id", anchor="center", width=70, stretch=False)
        t_cliente.column("razon", anchor="center", width=485, stretch=False)
        t_cliente.column("ruc", anchor="center", width=180, stretch=False)
        t_cliente.column("fecha", anchor="center", width=130, stretch=False)
        
        datos_clientes = [
            ("1", "Empresa ABC S.A.", "20123456789", "2023-11-01"),
            ("2", "Constructora XYZ", "10456789012", "2023-10-15"),
            ("3", "Servicios Industriales S.R.L.", "20345678901", "2023-09-22"),
            ("4", "Distribuidora Pérez", "10432109876", "2023-08-30"),
            ("5", "Inversiones López S.A.C.", "20567890123", "2023-07-19"),
            ("6", "Transportes Rojas", "10321098765", "2023-06-25"),
            ("7", "Tecnología Avanzada S.A.", "20198765432", "2023-05-14"),
            ("8", "Comercializadora García", "10456789023", "2023-04-07"),
            ("9", "Consultores Martínez S.A.", "20456789123", "2023-03-29"),
            ("10", "Agroindustria Chávez", "20543210987", "2023-02-18"),
            ("11", "Corporación Rivera", "10321987654", "2023-01-09"),
            ("12", "Soluciones Globales S.A.C.", "20345678901", "2022-12-11"),
            ("13", "Proyectos Vega", "10456789021", "2022-11-22"),
            ("14", "Servicios Integrales Torres", "20123456780", "2022-10-05"),
            ("15", "Constructora Lima S.A.C.", "20432109876", "2022-09-17"),
            ("16", "Empresa ABC S.A.", "20123456789", "2023-11-01"),
            ("17", "Constructora XYZ", "10456789012", "2023-10-15"),
            ("18", "Servicios Industriales S.R.L.", "20345678901", "2023-09-22"),
            ("19", "Distribuidora Pérez", "10432109876", "2023-08-30"),
            ("20", "Inversiones López S.A.C.", "20567890123", "2023-07-19"),
            ("21", "Consultores Martínez S.A.", "20456789123", "2023-03-29"),
        ]

        ejemplos_clientes = datos_clientes[:50]
        
        for cliente in ejemplos_clientes:
            t_cliente.insert("", "end", values=cliente)
            
        scrllbar_t_cli = ttk.Scrollbar(vent_clientes, orient="vertical", command=t_cliente.yview)
        t_cliente.configure(yscrollcommand=scrllbar_t_cli.set)
        scrllbar_t_cli.place(x=1177, y=80, height=569)
        
    def registrar_cliente(self):
        
        self.vent_clientes.withdraw()
        
        reg_cliente = tk.Toplevel(self.vent_clientes)
        reg_cliente.title("Registro de Clientes")
        reg_cliente.geometry("710x502")
        reg_cliente.resizable(False, False)
        reg_cliente.configure(bg="#373737")
        utils.centrar_ventana(reg_cliente)

        reg_cliente.protocol("WM_DELETE_WINDOW", lambda: None)
        
        canvas_reg_cli = tk.Canvas(reg_cliente, width=710, height=502, bg="#373737", highlightthickness=0)
        canvas_reg_cli.pack()
        
        utils.create_rounded_rectangle(canvas_reg_cli, 10, 10, 350, 136, radius=10, fill="#959595", outline="#959595")
        utils.create_rounded_rectangle(canvas_reg_cli, 10, 146, 350, 254, radius=10, fill="#959595", outline="#959595")
        #utils.create_rounded_rectangle(canvas_reg_cli, 10, 264, 350, 412, radius=10, fill="#959595", outline="#959595")
        utils.create_rounded_rectangle(canvas_reg_cli, 360, 10, 700, 118, radius=10, fill="#959595", outline="#959595")
        #utils.create_rounded_rectangle(canvas_reg_cli, 360, 128, 700, 246, radius=10, fill="#959595", outline="#959595")
        utils.create_rounded_rectangle(canvas_reg_cli, 360, 256, 700, 364, radius=10, fill="#959595", outline="#959595")
        #utils.create_rounded_rectangle(canvas_reg_cli, 360, 374, 700, 492, radius=10, fill="#959595", outline="#959595")
        
        canvas_reg_cli.create_text(20, 20, text="Razón Social / Empresa / Cliente", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_reg_cli, 20, 38, 340, 68, radius=10, fill="white", outline="#959595")
        reg_rs_cli = tk.Entry(reg_cliente, font=("Arial", 11), bd=0)
        reg_rs_cli.place(x=25, y=43, width=310, height=20)
        
        canvas_reg_cli.create_text(20, 78, text="RUC", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_reg_cli, 20, 96, 340, 126, radius=10, fill="white", outline="#959595")
        reg_ruc_cli = tk.Entry(reg_cliente, font=("Arial", 11), bd=0)
        reg_ruc_cli.place(x=25, y=101, width=310, height=20)
        
        canvas_reg_cli.create_text(20, 156, text="Persona de Contacto", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_reg_cli, 20, 174, 340, 204, radius=10, fill="white", outline="#959595")
        reg_persona = tk.Entry(reg_cliente, font=("Arial", 11), bd=0)
        reg_persona.place(x=25, y=179, width=310, height=20)
        
        canvas_reg_cli.create_text(370, 20, text="Área de Trabajo", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_reg_cli, 370, 38, 690, 68, radius=10, fill="white", outline="#959595")
        reg_ar_tb = tk.Entry(reg_cliente, font=("Arial", 11), bd=0)
        reg_ar_tb.place(x=375, y=43, width=310, height=20)

        canvas_reg_cli.create_text(370, 266, text="Dirección", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_reg_cli, 370, 284, 690, 314, radius=10, fill="white", outline="#959595")
        reg_direx = tk.Entry(reg_cliente, font=("Arial", 11), bd=0)
        reg_direx.place(x=375, y=289, width=310, height=20)
        
        btn_ag_persona = tk.Button(reg_cliente, text="Agregar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_ag_persona.place(x=20, y=214)
        
        btn_ed_persona = tk.Button(reg_cliente, text="Editar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=self.edit_persona_cont)
        btn_ed_persona.place(x=130, y=214)
        
        btn_del_persona = tk.Button(reg_cliente, text="Eliminar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_del_persona.place(x=240, y=214)
        
        
        btn_ag_trabajo = tk.Button(reg_cliente, text="Agregar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_ag_trabajo.place(x=370, y=78)
        
        btn_ed_trabajo = tk.Button(reg_cliente, text="Editar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=self.edit_area_trabajo)
        btn_ed_trabajo.place(x=480, y=78)
        
        btn_del_trabajo = tk.Button(reg_cliente, text="Eliminar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_del_trabajo.place(x=590, y=78)
        
        
        btn_ag_direx = tk.Button(reg_cliente, text="Agregar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_ag_direx.place(x=370, y=324)
        
        btn_ed_direx = tk.Button(reg_cliente, text="Editar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=self.edit_direx)
        btn_ed_direx.place(x=480, y=324)
        
        btn_del_direx = tk.Button(reg_cliente, text="Eliminar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_del_direx.place(x=590, y=324)
        
        
        btn_canc_reg = tk.Button(reg_cliente, text="Cancelar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_canc_reg.place(x=75, y=462)

        btn_gen = tk.Button(reg_cliente, text="Registrar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_gen.place(x=185, y=462)
        
        utils.aplicar_hover_a_botones([btn_ag_persona, btn_ed_persona, btn_del_persona, btn_ag_trabajo, btn_ed_trabajo, btn_del_trabajo, btn_ag_direx, btn_ed_direx, btn_del_direx, btn_canc_reg])
        
        t_persona = ttk.Treeview(reg_cliente, columns=("id_p", "persona"), show="headings", style="Custom.Treeview")
        t_persona.place(x=10, y=264, width=341, height=149)
        
        t_persona.heading("id_p", text="ID")
        t_persona.heading("persona", text="Persona de contacto")
        t_persona.column("id_p", anchor="center", width=35, stretch=False)
        t_persona.column("persona", anchor="center", width=290, stretch=False)
        
        datos_t_persona = [
            ("1", "Carlos Rodríguez"),
            ("2", "María López"),
            ("3", "Javier Martínez"),
            ("4", "Ana Fernández"),
            ("5", "Luis Gómez"),
        ]

        for persona in datos_t_persona:
            t_persona.insert("", "end", values=persona)
        
        scrollbar_t_persona = ttk.Scrollbar(reg_cliente, orient="vertical", command=t_persona.yview)
        t_persona.configure(yscrollcommand=scrollbar_t_persona.set)
        scrollbar_t_persona.place(x=337, y=264, height=149)
        
        
        t_area = ttk.Treeview(reg_cliente, columns=("id_a", "area"), show="headings", style="Custom.Treeview")
        t_area.place(x=360, y=128, width=341, height=119)
        
        t_area.heading("id_a", text="ID")
        t_area.heading("area", text="Área de Trabajo")
        t_area.column("id_a", anchor="center", width=35, stretch=False)
        t_area.column("area", anchor="center", width=290, stretch=False)
        
        datos_t_area = [
            ("1", "Recursos Humanos"),
            ("2", "Desarrollo de Software"),
            ("3", "Marketing"),
            ("4", "Finanzas"),
            ("5", "Logística")
        ]

        for dato in datos_t_area:
            t_area.insert("", "end", values=dato)
        
        scrollbar_t_area = ttk.Scrollbar(reg_cliente, orient="vertical", command=t_area.yview)
        t_area.configure(yscrollcommand=scrollbar_t_area.set)
        scrollbar_t_area.place(x=687, y=128, height=119)


        t_direx = ttk.Treeview(reg_cliente, columns=("id_d", "direx"), show="headings", style="Custom.Treeview")
        t_direx.place(x=360, y=374, width=341, height=119)
        
        t_direx.heading("id_d", text="ID")
        t_direx.heading("direx", text="Dirección")
        t_direx.column("id_d", anchor="center", width=35, stretch=False)
        t_direx.column("direx", anchor="center", width=290, stretch=False)
        
        datos_t_direx = [
            ("1", "Av. Los Olivos 123, Lima"),
            ("2", "Jr. San Martín 456, Arequipa"),
            ("3", "Calle Las Rosas 789, Cusco"),
            ("4", "Av. El Sol 321, Trujillo"),
            ("5", "Jr. Amazonas 654, Piura")
        ]
        
        for dato in datos_t_direx:
            t_direx.insert("", "end", values=dato)
        
        scrollbar_t_direx = ttk.Scrollbar(reg_cliente, orient="vertical", command=t_direx.yview)
        t_direx.configure(yscrollcommand=scrollbar_t_direx.set)
        scrollbar_t_direx.place(x=687, y=374, height=119)

    def edit_persona_cont(self):
        ed_pers = tk.Toplevel(self.vent_clientes)
        ed_pers.title("Editar Persona de Contacto")
        ed_pers.geometry("450x128")
        ed_pers.resizable(False, False)
        ed_pers.configure(bg="#373737")
        ed_pers.grab_set()
        utils.centrar_ventana(ed_pers)
        ed_pers.protocol("WM_DELETE_WINDOW", lambda: None)
        
        canvas_ed_pers = tk.Canvas(ed_pers, width=450, height=128, bg="#373737", highlightthickness=0)
        canvas_ed_pers.pack()
        
        utils.create_rounded_rectangle(canvas_ed_pers, 10, 10, 440, 78, radius=10, fill="#959595", outline="#959595")
        
        canvas_ed_pers.create_text(20, 20, text="Persona de Contacto", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_ed_pers, 20, 38, 430, 68, radius=10, fill="white", outline="#959595")
        inpt_ed_pers_con = tk.Entry(ed_pers, font=("Arial", 11), bd=0)
        inpt_ed_pers_con.place(x=25, y=43, width=400, height=20)
        
        btn_canc_pers = tk.Button(ed_pers, text="Cancelar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=ed_pers.destroy)
        btn_canc_pers.place(x=10, y=88)

        btn_save_pers = tk.Button(ed_pers, text="Guardar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_save_pers.place(x=120, y=88)
        
        utils.aplicar_hover_a_botones([btn_canc_pers, btn_save_pers])
        
    def edit_area_trabajo(self):
        ed_ar_trabajo = tk.Toplevel(self.vent_clientes)
        ed_ar_trabajo.title("Editar Área de Trabajo")
        ed_ar_trabajo.geometry("450x128")
        ed_ar_trabajo.resizable(False, False)
        ed_ar_trabajo.configure(bg="#373737")
        ed_ar_trabajo.grab_set()
        utils.centrar_ventana(ed_ar_trabajo)
        ed_ar_trabajo.protocol("WM_DELETE_WINDOW", lambda: None)
        
        canvas_ed_ar_trab = tk.Canvas(ed_ar_trabajo, width=450, height=128, bg="#373737", highlightthickness=0)
        canvas_ed_ar_trab.pack()

        utils.create_rounded_rectangle(canvas_ed_ar_trab, 10, 10, 440, 78, radius=10, fill="#959595", outline="#959595")
    
        canvas_ed_ar_trab.create_text(20, 20, text="Área de Trabajo", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_ed_ar_trab, 20, 38, 430, 68, radius=10, fill="white", outline="#959595")
        inpt_ed_ar_trab = tk.Entry(ed_ar_trabajo, font=("Arial", 11), bd=0)
        inpt_ed_ar_trab.place(x=25, y=43, width=400, height=20)
        
        btn_canc_ar_trab = tk.Button(ed_ar_trabajo, text="Cancelar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=ed_ar_trabajo.destroy)
        btn_canc_ar_trab.place(x=10, y=88)

        btn_save_ar_trab = tk.Button(ed_ar_trabajo, text="Guardar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_save_ar_trab.place(x=120, y=88)
        
        utils.aplicar_hover_a_botones([btn_canc_ar_trab, btn_save_ar_trab])
        
    def edit_direx(self):
        ed_direccion = tk.Toplevel(self.vent_clientes)
        ed_direccion.title("Editar Dirección")
        ed_direccion.geometry("450x128")
        ed_direccion.resizable(False, False)
        ed_direccion.configure(bg="#373737")
        ed_direccion.grab_set()
        utils.centrar_ventana(ed_direccion)
        ed_direccion.protocol("WM_DELETE_WINDOW", lambda: None)
        
        canvas_ed_direccion = tk.Canvas(ed_direccion, width=450, height=128, bg="#373737", highlightthickness=0)
        canvas_ed_direccion.pack()

        utils.create_rounded_rectangle(canvas_ed_direccion, 10, 10, 440, 78, radius=10, fill="#959595", outline="#959595")

        canvas_ed_direccion.create_text(20, 20, text="Dirección", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_ed_direccion, 20, 38, 430, 68, radius=10, fill="white", outline="#959595")
        inpt_ed_direx = tk.Entry(ed_direccion, font=("Arial", 11), bd=0)
        inpt_ed_direx.place(x=25, y=43, width=400, height=20)

        btn_canc_direx = tk.Button(ed_direccion, text="Cancelar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=ed_direccion.destroy)
        btn_canc_direx.place(x=10, y=88)

        btn_save_direx = tk.Button(ed_direccion, text="Guardar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_save_direx.place(x=120, y=88)
        
        utils.aplicar_hover_a_botones([btn_canc_direx, btn_save_direx])

    def detalle_cliente(self):
        
        self.vent_clientes.withdraw()
        
        det_cliente = tk.Toplevel(self.vent_clientes)
        det_cliente.title("Detalle del Cliente")
        det_cliente.geometry("710x502")
        det_cliente.resizable(False, False)
        det_cliente.configure(bg="#373737")
        utils.centrar_ventana(det_cliente)
        
        det_cliente.protocol("WM_DELETE_WINDOW", lambda: None)

        canvas_det_cli = tk.Canvas(det_cliente, width=710, height=502, bg="#373737", highlightthickness=0)
        canvas_det_cli.pack()
        
        utils.create_rounded_rectangle(canvas_det_cli, 10, 10, 350, 136, radius=10, fill="#959595", outline="#959595")
        utils.create_rounded_rectangle(canvas_det_cli, 10, 146, 350, 254, radius=10, fill="#959595", outline="#959595")
        #utils.create_rounded_rectangle(canvas_det_cli, 10, 264, 350, 412, radius=10, fill="#959595", outline="#959595")
        utils.create_rounded_rectangle(canvas_det_cli, 360, 10, 700, 118, radius=10, fill="#959595", outline="#959595")
        #utils.create_rounded_rectangle(canvas_det_cli, 360, 128, 700, 246, radius=10, fill="#959595", outline="#959595")
        utils.create_rounded_rectangle(canvas_det_cli, 360, 256, 700, 364, radius=10, fill="#959595", outline="#959595")
        #utils.create_rounded_rectangle(canvas_det_cli, 360, 374, 700, 492, radius=10, fill="#959595", outline="#959595")

        canvas_det_cli.create_text(20, 20, text="Razón Social / Empresa / Cliente", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_det_cli, 20, 38, 340, 68, radius=10, fill="white", outline="#959595")
        det_rs_cli = tk.Entry(det_cliente, font=("Arial", 11), bd=0)
        det_rs_cli.place(x=25, y=43, width=310, height=20)
        
        canvas_det_cli.create_text(20, 78, text="RUC", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_det_cli, 20, 96, 340, 126, radius=10, fill="white", outline="#959595")
        det_ruc_cli = tk.Entry(det_cliente, font=("Arial", 11), bd=0)
        det_ruc_cli.place(x=25, y=101, width=310, height=20)
        
        canvas_det_cli.create_text(20, 156, text="Persona de Contacto", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_det_cli, 20, 174, 340, 204, radius=10, fill="white", outline="#959595")
        det_persona = tk.Entry(det_cliente, font=("Arial", 11), bd=0)
        det_persona.place(x=25, y=179, width=310, height=20)
        
        canvas_det_cli.create_text(370, 20, text="Área de Trabajo", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_det_cli, 370, 38, 690, 68, radius=10, fill="white", outline="#959595")
        det_ar_tb = tk.Entry(det_cliente, font=("Arial", 11), bd=0)
        det_ar_tb.place(x=375, y=43, width=310, height=20)

        canvas_det_cli.create_text(370, 266, text="Dirección", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_det_cli, 370, 284, 690, 314, radius=10, fill="white", outline="#959595")
        det_direx = tk.Entry(det_cliente, font=("Arial", 11), bd=0)
        det_direx.place(x=375, y=289, width=310, height=20)
        
        btn_det_ag_persona = tk.Button(det_cliente, text="Agregar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_det_ag_persona.place(x=20, y=214)
        
        btn_det_ed_persona = tk.Button(det_cliente, text="Editar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=self.edit_persona_cont)
        btn_det_ed_persona.place(x=130, y=214)
        
        btn_det_del_persona = tk.Button(det_cliente, text="Eliminar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_det_del_persona.place(x=240, y=214)
        
        
        btn_det_ag_trabajo = tk.Button(det_cliente, text="Agregar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_det_ag_trabajo.place(x=370, y=78)
        
        btn_det_ed_trabajo = tk.Button(det_cliente, text="Editar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=self.edit_area_trabajo)
        btn_det_ed_trabajo.place(x=480, y=78)
        
        btn_det_del_trabajo = tk.Button(det_cliente, text="Eliminar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_det_del_trabajo.place(x=590, y=78)
        
        
        btn_det_ag_direx = tk.Button(det_cliente, text="Agregar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_det_ag_direx.place(x=370, y=324)
        
        btn_det_ed_direx = tk.Button(det_cliente, text="Editar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=self.edit_direx)
        btn_det_ed_direx.place(x=480, y=324)
        
        btn_det_del_direx = tk.Button(det_cliente, text="Eliminar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_det_del_direx.place(x=590, y=324)
        
        
        btn_det_canc = tk.Button(det_cliente, text="Cancelar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_det_canc.place(x=75, y=462)

        btn_det_save = tk.Button(det_cliente, text="Guardar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_det_save.place(x=185, y=462)
        
        utils.aplicar_hover_a_botones([btn_det_ag_persona, btn_det_ed_persona, btn_det_del_persona, btn_det_ag_trabajo, 
                                       btn_det_ed_trabajo, btn_det_del_trabajo, btn_det_ag_direx, btn_det_ed_direx, 
                                       btn_det_del_direx, btn_det_canc, btn_det_save])
        
        t_det_persona = ttk.Treeview(det_cliente, columns=("id_dp", "dpersona"), show="headings", style="Custom.Treeview")
        t_det_persona.place(x=10, y=264, width=341, height=149)
        
        t_det_persona.heading("id_dp", text="ID")
        t_det_persona.heading("dpersona", text="Persona de contacto")
        t_det_persona.column("id_dp", anchor="center", width=35, stretch=False)
        t_det_persona.column("dpersona", anchor="center", width=290, stretch=False)
        
        datos_t_dpersona = [
            ("1", "AAA"),
            ("2", "BBB"),
            ("3", "CCC"),
            ("4", "DDD"),
            ("5", "EEE"),
        ]

        for d_persona in datos_t_dpersona:
            t_det_persona.insert("", "end", values=d_persona)
        
        scrollbar_t_dpersona = ttk.Scrollbar(det_cliente, orient="vertical", command=t_det_persona.yview)
        t_det_persona.configure(yscrollcommand=scrollbar_t_dpersona.set)
        scrollbar_t_dpersona.place(x=337, y=264, height=149)
        
        
        t_det_area = ttk.Treeview(det_cliente, columns=("id_da", "darea"), show="headings", style="Custom.Treeview")
        t_det_area.place(x=360, y=128, width=341, height=119)
        
        t_det_area.heading("id_da", text="ID")
        t_det_area.heading("darea", text="Área de Trabajo")
        t_det_area.column("id_da", anchor="center", width=35, stretch=False)
        t_det_area.column("darea", anchor="center", width=290, stretch=False)
        
        datos_t_darea = [
            ("1", "AAA"),
            ("2", "BBB"),
            ("3", "CCC"),
            ("4", "DDD"),
            ("5", "EEE")
        ]

        for d_area in datos_t_darea:
            t_det_area.insert("", "end", values=d_area)
        
        scrollbar_t_darea = ttk.Scrollbar(det_cliente, orient="vertical", command=t_det_area.yview)
        t_det_area.configure(yscrollcommand=scrollbar_t_darea.set)
        scrollbar_t_darea.place(x=687, y=128, height=119)
        
        
        t_det_direx = ttk.Treeview(det_cliente, columns=("id_dd", "ddirex"), show="headings", style="Custom.Treeview")
        t_det_direx.place(x=360, y=374, width=341, height=119)
        
        t_det_direx.heading("id_dd", text="ID")
        t_det_direx.heading("ddirex", text="Dirección")
        t_det_direx.column("id_dd", anchor="center", width=35, stretch=False)
        t_det_direx.column("ddirex", anchor="center", width=290, stretch=False)
        
        datos_t_ddirex = [
            ("1", "AAA"),
            ("2", "BBB"),
            ("3", "CCC"),
            ("4", "DDD"),
            ("5", "EEE")
        ]
        
        for d_direx in datos_t_ddirex:
            t_det_direx.insert("", "end", values=d_direx)
        
        scrollbar_t_ddirex = ttk.Scrollbar(det_cliente, orient="vertical", command=t_det_direx.yview)
        t_det_direx.configure(yscrollcommand=scrollbar_t_ddirex.set)
        scrollbar_t_ddirex.place(x=687, y=374, height=119)   


class cotizaciones:
    def __init__(self, root, vent_coti):
        
        self.root = root
        self.vent_coti = vent_coti
        self.root.withdraw()
        
        self.vent_coti = vent_coti
        self.vent_coti.title("Registro de de Cotizaciones")
        self.vent_coti.geometry("1300x608")
        self.vent_coti.resizable(False, False)
        self.vent_coti.configure(bg="#373737")
        utils.centrar_ventana(self.vent_coti)
        
        self.alerta = alertas(vent_coti)
        
        self.vent_coti.protocol("WM_DELETE_WINDOW", lambda: None)
        
        canvas_cotizacion = tk.Canvas(vent_coti, width=1300, height=608, bg="#373737", highlightthickness=0)
        canvas_cotizacion.pack()
        
        utils.create_rounded_rectangle(canvas_cotizacion, 10, 10, 300, 410, radius=10, fill="#959595", outline="#959595")
        utils.create_rounded_rectangle(canvas_cotizacion, 10, 420, 300, 598, radius=10, fill="#959595", outline="#959595")
        #utils.create_rounded_rectangle(canvas_cotizacion, 310, 80, 1290, 558, radius=10, fill="#959595", outline="#959595")
        
        canvas_cotizacion.create_text(435, 20, text="REGISTRO DE COTIZACIONES", anchor="nw", font=("Raleway", 20, "bold"), fill="White")
        canvas_cotizacion.create_text(20, 22, text="Opciones", anchor="nw", font=("Raleway", 20, "bold"), fill="White")
        canvas_cotizacion.create_text(20, 432, text="Filtros", anchor="nw", font=("Raleway", 20, "bold"), fill="White")
        
        btn_reg_coti = tk.Button(vent_coti, text="Registrar Cotización", width=37, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=self.registrar_cotizacion)
        btn_reg_coti.place(x=22, y=80)
        
        btn_seg_cot = tk.Button(vent_coti, text="Seguimiento / Detalles", width=37, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=self.seguimiento_cotizacion)
        btn_seg_cot.place(x=22, y=125)
        
        btn_op_cot = tk.Button(vent_coti, text="Abrir Carpeta", width=37, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_op_cot.place(x=22, y=170)
        
        btn_menu_cot = tk.Button(vent_coti, text="Volver al inicio", width=37, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_menu_cot.place(x=22, y=215)
        
        btn_sig_cot = tk.Button(vent_coti, text="Siguiente", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_sig_cot.place(x=1190, y=568)
        
        btn_atras_cot = tk.Button(vent_coti, text="Anterior", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_atras_cot.place(x=1080, y=568)
        
        utils.aplicar_hover_a_botones([btn_reg_coti, btn_seg_cot, btn_op_cot, btn_menu_cot, btn_sig_cot, btn_atras_cot])
        
        canvas_cotizacion.create_text(20, 486, text="Por Estado", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        
        cbo_estado = ttk.Combobox(vent_coti, values=["Todos los registros", "Aprobado", "Pendiente", "No Aprobado"], state="readonly", font=("Raleway", 10))
        cbo_estado.place(x=20, y=504, width=270, height=30)
        cbo_estado.current(0)
        
        canvas_cotizacion.create_text(20, 540, text="Por Fecha", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        
        cbo_date_cot = ttk.Combobox(vent_coti, values=["Todos los registros", "Última semana", "Últimas 2 semanas", "Último mes", "Últimos 3 meses", "Últimos 6 meses", "Último año"], state="readonly", font=("Raleway", 10))
        cbo_date_cot.place(x=20, y=558, width=270, height=30)
        cbo_date_cot.current(0)
        
        cbo_page_cot = ttk.Combobox(vent_coti, values=["1", "2"], state="readonly", font=("Arial", 10))
        cbo_page_cot.place(x=310, y=568, width=70, height=30)
        cbo_page_cot.current(0)

        search_canvas_cot = tk.Canvas(vent_coti, width=350, height=40, bg="#373737", highlightthickness=0)
        search_canvas_cot.place(x=940, y=20)
        
        utils.create_rounded_rectangle(search_canvas_cot, 0, 0, 350, 40, radius=10, fill="white", outline="gray")
        search_canvas_cot.create_line(315, 7, 315, 34, fill="gray", width=2)
        
        search_icon_path = os.path.join(ICON_DIR, "search.png")
        try:
            search_icon_cot = tk.PhotoImage(file=search_icon_path)
            search_icon_id_cot = search_canvas_cot.create_image(321, 8, anchor="nw", image=search_icon_cot)
            search_canvas_cot.image = search_icon_cot
        except Exception as e:
            raise FileNotFoundError(f"El archivo del icono no se encontró en la ruta: {search_icon_path}. Error: {e}")

        search_canvas_cot.tag_bind(search_icon_id_cot, "<Button-1>", lambda e: self.alerta.no_datos())
        
        search_entry_cot = tk.Entry(search_canvas_cot, font=("Arial", 13), width=40, bd=0, relief="flat", fg='grey')
        search_entry_cot.insert(0, "Buscar...")
        search_entry_cot.bind("<FocusIn>", lambda event: utils.clear_placeholder(event, search_entry_cot))
        search_entry_cot.bind("<FocusOut>", lambda event: utils.placeholder_search(event, search_entry_cot))
        search_entry_cot.place(x=6, y=7, width=301, height=27)

        t_cotizacion = ttk.Treeview(vent_coti, columns=("id_c", "cliente_c", "servicio_c", "estado_c", "fecha_c"), show="headings", style="Custom.Treeview")
        t_cotizacion.place(x=310, y=80, width=981, height=479)
        
        t_cotizacion.heading("id_c", text="Nro")
        t_cotizacion.heading("cliente_c", text="Cliente / Empresa")
        t_cotizacion.heading("servicio_c", text="Servicio")
        t_cotizacion.heading("estado_c", text="Estado")
        t_cotizacion.heading("fecha_c", text="Fecha")
        
        t_cotizacion.column("id_c", anchor="center", width=100, stretch=False)
        t_cotizacion.column("cliente_c", anchor="center", width=268, stretch=False)
        t_cotizacion.column("servicio_c", anchor="center", width=387, stretch=False)
        t_cotizacion.column("estado_c", anchor="center", width=110, stretch=False)
        t_cotizacion.column("fecha_c", anchor="center", width=100, stretch=False)
        
        datos_t_cotizacion = [
            ("001/1000", "Empresa A", "Servicio de Mantenimiento", "Pendiente", "2023-11-01"),
            ("2", "Empresa B", "Desarrollo de Software", "Aceptado", "2023-10-15"),
            ("3", "Empresa C", "Consultoría en TI", "Rechazado", "2023-09-20"),
            ("4", "Empresa D", "Instalación de Redes", "Pendiente", "2023-08-25"),
            ("5", "Empresa E", "Auditoría de Sistemas", "Aceptado", "2023-07-18"),
            ("6", "Empresa F", "Migración de Datos", "Pendiente", "2023-06-22"),
            ("7", "Empresa G", "Optimización de Procesos", "Aceptado", "2023-05-30"),
            ("8", "Empresa H", "Soporte Técnico", "Rechazado", "2023-04-12"),
            ("9", "Empresa I", "Desarrollo Web", "Pendiente", "2023-03-28"),
            ("10", "Empresa J", "Evaluación de Seguridad", "Aceptado", "2023-02-11"),
            ("11", "Empresa K", "Diseño Gráfico", "Pendiente", "2023-01-15"),
            ("12", "Empresa L", "Análisis de Datos", "Aceptado", "2022-12-05"),
            ("13", "Empresa M", "Capacitación de Personal", "Rechazado", "2022-11-22"),
            ("14", "Empresa N", "Optimización de Redes", "Pendiente", "2022-10-13"),
            ("15", "Empresa O", "Desarrollo de App Móvil", "Aceptado", "2022-09-07"),
            ("16", "Empresa P", "Soporte en Ciberseguridad", "Pendiente", "2022-08-20"),
            ("17", "Empresa Q", "Automatización de Procesos", "Aceptado", "2022-07-11"),
            ("18", "Empresa R", "Asesoría en Transformación Digital", "Rechazado", "2022-06-02"),
            ("19", "Empresa S", "Administración de Sistemas", "Pendiente", "2022-05-19"),
            ("20", "Empresa T", "Desarrollo de Contenido", "Aceptado", "2022-04-25"),
        ]
        
        ejemplo_cotzacion = datos_t_cotizacion[:50]
        
        for cotizacion in ejemplo_cotzacion:
            t_cotizacion.insert("", "end", values=cotizacion)
        
        scrllbar_t_coti = ttk.Scrollbar(vent_coti, orient="vertical", command=t_cotizacion.yview)
        t_cotizacion.configure(yscrollcommand=scrllbar_t_coti.set)
        scrllbar_t_coti.place(x=1277, y=80, height=479)
        
    def registrar_cotizacion(self):
        
        self.vent_coti.withdraw()
        
        reg_coti = tk.Toplevel(self.vent_coti)
        reg_coti.title("Registro de Cotización")
        reg_coti.geometry("590x380")
        reg_coti.resizable(False, False)
        reg_coti.configure(bg="#373737")
        utils.centrar_ventana(reg_coti)
        
        reg_coti.protocol("WM_DELETE_WINDOW", lambda: None)
        
        canvas_reg_coti = tk.Canvas(reg_coti, width=590, height=380, bg="#373737", highlightthickness=0)
        canvas_reg_coti.pack()
        
        utils.create_rounded_rectangle(canvas_reg_coti, 10, 10, 580, 252, radius=10, fill="#959595", outline="#959595")
        utils.create_rounded_rectangle(canvas_reg_coti, 10, 262, 580, 330, radius=10, fill="#959595", outline="#959595")
        
        canvas_reg_coti.create_text(20, 20, text="Nro de Cotización", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_reg_coti, 20, 38, 170, 68, radius=10, fill="white", outline="#959595")
        reg_cot = tk.Entry(reg_coti, font=("Arial", 11), bd=0)
        reg_cot.place(x=25, y=43, width=140, height=20)
        
        canvas_reg_coti.create_text(180, 20, text="Cliente / Empresa", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_reg_coti, 180, 38, 380, 68, radius=10, fill="white", outline="#959595")
        reg_cli_cot = tk.Entry(reg_coti, font=("Arial", 11), bd=0)
        reg_cli_cot.place(x=185, y=43, width=190, height=20)
        
        canvas_reg_coti.create_text(390, 20, text="Área de Trabajo", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        cbo_area_cot = ttk.Combobox(reg_coti, values=["Escoja una Opción", "Ejemplo1"], state="readonly", font=("Raleway", 10))
        cbo_area_cot.place(x=390, y=38, width=180, height=31)
        cbo_area_cot.current(0)
        
        canvas_reg_coti.create_text(20, 78, text="Persona de Contacto", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        cbo_per_cot = ttk.Combobox(reg_coti, values=["Escoja una Opción", "Ejemplo1"], state="readonly", font=("Raleway", 10))
        cbo_per_cot.place(x=20, y=96, width=200, height=31)
        cbo_per_cot.current(0)
        
        canvas_reg_coti.create_text(230, 78, text="Título del Servicio", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_reg_coti, 230, 96, 570, 126, radius=10, fill="white", outline="#959595")
        reg_serv = tk.Entry(reg_coti, font=("Arial", 11), bd=0)
        reg_serv.place(x=235, y=101, width=330, height=20)
        
        canvas_reg_coti.create_text(20, 136, text="Tiempo de Ejecución", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_reg_coti, 20, 154, 250, 184, radius=10, fill="white", outline="#959595")
        reg_tiempo = tk.Entry(reg_coti, font=("Arial", 11), bd=0)
        reg_tiempo.place(x=25, y=159, width=220, height=20)
        
        canvas_reg_coti.create_text(260, 136, text="Forma de Pago", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_reg_coti, 260, 154, 430, 184, radius=10, fill="white", outline="#959595")
        reg_pago = tk.Entry(reg_coti, font=("Arial", 11), bd=0)
        reg_pago.place(x=265, y=159, width=160, height=20)
        
        canvas_reg_coti.create_text(440, 136, text="Costo Total", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_reg_coti, 440, 154, 570, 184, radius=10, fill="white", outline="#959595")
        reg_costo = tk.Entry(reg_coti, font=("Arial", 11), bd=0)
        reg_costo.place(x=445, y=159, width=120, height=20)
        
        canvas_reg_coti.create_text(20, 193, text="Fecha", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        date_coti = DateEntry(reg_coti, font=("Raleway", 11),state="readonly" , width=17, background='#4A6984', foreground='white', borderwidth=1)
        date_coti.place(x=20, y=211, width=150, height=30)
        
        canvas_reg_coti.create_text(180, 194, text="Estado", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        cbo_reg_est = ttk.Combobox(reg_coti, values=["Escoja una Opción", "Aprobado", "Pendiente", "No Aprobado"], state="readonly", font=("Raleway", 10))
        cbo_reg_est.place(x=180, y=212, width=150, height=31)
        cbo_reg_est.current(0)
        
        canvas_reg_coti.create_text(20, 272, text="Cotización", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        
        btn_ad_cot = tk.Button(reg_coti, text="Adjuntar", command=lambda: utils.adjuntar_archivo(label_coti, "cotizacion"), font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_ad_cot.place(x=20, y=290, width=140, height=30)
        
        label_coti = tk.Label(reg_coti, text="Cotización", font=("Raleway", 9), bg="#373737", fg="white")
        label_coti.place(x=170, y=290, width=310, height=30)
        
        btn_ver_cot = tk.Button(reg_coti, text="Abrir", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_ver_cot.place(x=490, y=290, width=80, height=30)
        
        btn_canc_cot = tk.Button(reg_coti, text="Cancelar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_canc_cot.place(x=190, y=340)
        
        btn_reg_cotiz = tk.Button(reg_coti, text="Registrar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_reg_cotiz.place(x=300, y=340)
        
        utils.aplicar_hover_a_botones([btn_ad_cot, btn_ver_cot, btn_canc_cot, btn_reg_cotiz])
        
    def seguimiento_cotizacion(self):
        
        self.vent_coti.withdraw()
        
        seg_coti = tk.Toplevel(self.vent_coti)
        seg_coti.title("Seguimiento de Cotización")
        seg_coti.geometry("590x380")
        seg_coti.resizable(False, False)
        seg_coti.configure(bg="#373737")
        utils.centrar_ventana(seg_coti)
        seg_coti.protocol("WM_DELETE_WINDOW", lambda: None)
        
        canvas_det_cot = tk.Canvas(seg_coti, width=590, height=380, bg="#373737", highlightthickness=0)
        canvas_det_cot.pack()
        
        utils.create_rounded_rectangle(canvas_det_cot, 10, 10, 580, 252, radius=10, fill="#959595", outline="#959595")
        utils.create_rounded_rectangle(canvas_det_cot, 10, 262, 580, 330, radius=10, fill="#959595", outline="#959595")
        
        canvas_det_cot.create_text(20, 20, text="Nro de Cotización", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_det_cot, 20, 38, 170, 68, radius=10, fill="white", outline="#959595")
        det_cot = tk.Entry(seg_coti, font=("Arial", 11), bd=0)
        det_cot.place(x=25, y=43, width=140, height=20)
        
        canvas_det_cot.create_text(180, 20, text="Cliente / Empresa", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_det_cot, 180, 38, 380, 68, radius=10, fill="white", outline="#959595")
        det_cli_cot = tk.Entry(seg_coti, font=("Arial", 11), bd=0)
        det_cli_cot.place(x=185, y=43, width=190, height=20)
        
        canvas_det_cot.create_text(390, 20, text="Área de Trabajo", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        cbo_dar_cot = ttk.Combobox(seg_coti, values=["Escoja una Opción", "Ejemplo1"], state="readonly", font=("Raleway", 10))
        cbo_dar_cot.place(x=390, y=38, width=180, height=31)
        cbo_dar_cot.current(0)
        
        canvas_det_cot.create_text(20, 78, text="Persona de Contacto", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        cbo_dper_cot = ttk.Combobox(seg_coti, values=["Escoja una Opción", "Ejemplo1"], state="readonly", font=("Raleway", 10))
        cbo_dper_cot.place(x=20, y=96, width=200, height=31)
        cbo_dper_cot.current(0)
        
        canvas_det_cot.create_text(230, 78, text="Título del Servicio", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_det_cot, 230, 96, 570, 126, radius=10, fill="white", outline="#959595")
        det_serv = tk.Entry(seg_coti, font=("Arial", 11), bd=0)
        det_serv.place(x=235, y=101, width=330, height=20)
        
        canvas_det_cot.create_text(20, 136, text="Tiempo de Ejecución", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_det_cot, 20, 154, 250, 184, radius=10, fill="white", outline="#959595")
        det_tiempo = tk.Entry(seg_coti, font=("Arial", 11), bd=0)
        det_tiempo.place(x=25, y=159, width=220, height=20)
        
        canvas_det_cot.create_text(260, 136, text="Forma de Pago", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_det_cot, 260, 154, 430, 184, radius=10, fill="white", outline="#959595")
        det_pago = tk.Entry(seg_coti, font=("Arial", 11), bd=0)
        det_pago.place(x=265, y=159, width=160, height=20)
        
        canvas_det_cot.create_text(440, 136, text="Costo Total", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_det_cot, 440, 154, 570, 184, radius=10, fill="white", outline="#959595")
        det_costo = tk.Entry(seg_coti, font=("Arial", 11), bd=0)
        det_costo.place(x=445, y=159, width=120, height=20)
        
        canvas_det_cot.create_text(20, 193, text="Fecha", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        date_seg_coti = DateEntry(seg_coti, font=("Raleway", 11),state="readonly" , width=17, background='#4A6984', foreground='white', borderwidth=1)
        date_seg_coti.place(x=20, y=211, width=150, height=30)
        
        canvas_det_cot.create_text(180, 194, text="Estado", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        cbo_det_est = ttk.Combobox(seg_coti, values=["Escoja una Opción", "Aprobado", "Pendiente", "No Aprobado"], state="readonly", font=("Raleway", 10))
        cbo_det_est.place(x=180, y=212, width=150, height=31)
        cbo_det_est.current(0)
        
        canvas_det_cot.create_text(20, 272, text="Cotización", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        
        btn_det_cot = tk.Button(seg_coti, text="Cambiar documento", command=lambda: utils.adjuntar_archivo(label_seg_coti, "cotizacion"), font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_det_cot.place(x=20, y=290, width=140, height=30)
        
        label_seg_coti = tk.Label(seg_coti, text="Cotización", font=("Raleway", 9), bg="#373737", fg="white")
        label_seg_coti.place(x=170, y=290, width=310, height=30)
        
        btn_ver_det_cot = tk.Button(seg_coti, text="Abrir", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_ver_det_cot.place(x=490, y=290, width=80, height=30)
        
        btn_canc_dcot = tk.Button(seg_coti, text="Cancelar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_canc_dcot.place(x=190, y=340)
        
        btn_reg_dcot = tk.Button(seg_coti, text="Guardar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_reg_dcot.place(x=300, y=340)
        
        utils.aplicar_hover_a_botones([btn_det_cot, btn_ver_det_cot, btn_canc_dcot, btn_reg_dcot])


class orden_compra:
    def __init__ (self, root, vent_oc):
        
        self.root = root
        self.vent_oc = vent_oc
        self.root.withdraw()
        
        self.vent_oc = vent_oc
        self.vent_oc.title("Registro de Orden de Compra")
        self.vent_oc.geometry("1300x608")
        self.vent_oc.resizable(False, False)
        self.vent_oc.configure(bg="#373737")
        utils.centrar_ventana(self.vent_oc)
        
        self.alerta = alertas(vent_oc)
        
        self.vent_oc.protocol("WM_DELETE_WINDOW", lambda: None)
        
        canvas_oc = tk.Canvas(vent_oc, width=1300, height=608, bg="#373737", highlightthickness=0)
        canvas_oc.pack()
        
        utils.create_rounded_rectangle(canvas_oc, 10, 10, 300, 410, radius=10, fill="#959595", outline="#959595")
        utils.create_rounded_rectangle(canvas_oc, 10, 420, 300, 598, radius=10, fill="#959595", outline="#959595")
        #utils.create_rounded_rectangle(canvas_oc, 310, 80, 1290, 558, radius=10, fill="#959595", outline="#959595")
        
        canvas_oc.create_text(400, 20, text="REGISTRO DE ORDEN DE COMPRA", anchor="nw", font=("Raleway", 20, "bold"), fill="White")
        canvas_oc.create_text(20, 22, text="Opciones", anchor="nw", font=("Raleway", 20, "bold"), fill="White")
        canvas_oc.create_text(20, 432, text="Filtros", anchor="nw", font=("Raleway", 20, "bold"), fill="White")
        
        btn_reg_ocompra = tk.Button(vent_oc, text="Registrar Orden de Compra", width=37, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=self.registrar_orden_compra)
        btn_reg_ocompra.place(x=22, y=80)
        
        btn_seg_oc = tk.Button(vent_oc, text="Seguimiento / Detalles", width=37, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=self.seguimiento_orden_compra)
        btn_seg_oc.place(x=22, y=125)
        
        btn_op_oc = tk.Button(vent_oc, text="Abrir Carpeta", width=37, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_op_oc.place(x=22, y=170)
        
        btn_menu_oc = tk.Button(vent_oc, text="Volver al inicio", width=37, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_menu_oc.place(x=22, y=215)
        
        btn_sig_oc = tk.Button(vent_oc, text="Siguiente", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_sig_oc.place(x=1190, y=568)
        
        btn_atras_oc = tk.Button(vent_oc, text="Anterior", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_atras_oc.place(x=1080, y=568)
        
        utils.aplicar_hover_a_botones([btn_reg_ocompra, btn_seg_oc, btn_op_oc, btn_menu_oc, btn_sig_oc, btn_atras_oc])
        
        canvas_oc.create_text(20, 486, text="Por Estado", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        
        cbo_estado_oc = ttk.Combobox(vent_oc, values=["Todos los registros", "Trabajo No Iniciado", "En Proceso", "Completado Parcialmente", "Finalizado"], state="readonly", font=("Raleway", 10))
        cbo_estado_oc.place(x=20, y=504, width=270, height=30)
        cbo_estado_oc.current(0)
        
        canvas_oc.create_text(20, 540, text="Por Fecha", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        
        cbo_date_oc = ttk.Combobox(vent_oc, values=["Todos los registros", "Última semana", "Últimas 2 semanas", "Último mes", "Últimos 3 meses", "Últimos 6 meses", "Último año"], state="readonly", font=("Raleway", 10))
        cbo_date_oc.place(x=20, y=558, width=270, height=30)
        cbo_date_oc.current(0)
        
        cbo_page_oc = ttk.Combobox(vent_oc, values=["1", "2"], state="readonly", font=("Arial", 10))
        cbo_page_oc.place(x=310, y=568, width=70, height=30)
        cbo_page_oc.current(0)
        
        search_canvas_oc = tk.Canvas(vent_oc, width=350, height=40, bg="#373737", highlightthickness=0)
        search_canvas_oc.place(x=940, y=20)
        
        utils.create_rounded_rectangle(search_canvas_oc, 0, 0, 350, 40, radius=10, fill="white", outline="gray")
        search_canvas_oc.create_line(315, 7, 315, 34, fill="gray", width=2)
        
        search_icon_path = os.path.join(ICON_DIR, "search.png")
        try:
            search_icon_oc = tk.PhotoImage(file=search_icon_path)
            search_icon_id_oc = search_canvas_oc.create_image(321, 8, anchor="nw", image=search_icon_oc)
            search_canvas_oc.image = search_icon_oc
        except Exception as e:
            raise FileNotFoundError(f"El archivo del icono no se encontró en la ruta: {search_icon_path}. Error: {e}")

        search_canvas_oc.tag_bind(search_icon_id_oc, "<Button-1>", lambda e: self.alerta.no_datos())
        
        search_entry_oc = tk.Entry(search_canvas_oc, font=("Arial", 13), width=40, bd=0, relief="flat", fg='grey')
        search_entry_oc.insert(0, "Buscar...")
        search_entry_oc.bind("<FocusIn>", lambda event: utils.clear_placeholder(event, search_entry_oc))
        search_entry_oc.bind("<FocusOut>", lambda event: utils.placeholder_search(event, search_entry_oc))
        search_entry_oc.place(x=6, y=7, width=301, height=27)
        
        t_orden = ttk.Treeview(vent_oc, columns=("id_oc", "cliente_oc", "servicio_oc", "estado_oc", "fecha_oc"), show="headings", style="Custom.Treeview")
        t_orden.place(x=310, y=80, width=981, height=479)
        
        t_orden.heading("id_oc", text="Nro")
        t_orden.heading("cliente_oc", text="Cliente / Empresa")
        t_orden.heading("servicio_oc", text="Servicio")
        t_orden.heading("estado_oc", text="Estado")
        t_orden.heading("fecha_oc", text="Fecha")
        
        t_orden.column("id_oc", anchor="center", width=100, stretch=False)
        t_orden.column("cliente_oc", anchor="center", width=188, stretch=False)
        t_orden.column("servicio_oc", anchor="center", width=387, stretch=False)
        t_orden.column("estado_oc", anchor="center", width=190, stretch=False)
        t_orden.column("fecha_oc", anchor="center", width=100, stretch=False)
        
        datos_t_orden = [
            ("1", "Empresa A", "Servicio de Mantenimiento", "Trabajo No Iniciado", "2023-11-01"),
            ("2", "Empresa B", "Desarrollo de Software", "En Proceso", "2023-10-15"),
            ("3", "Empresa C", "Consultoría en TI", "Completado Parcialmente", "2023-09-20"),
            ("4", "Empresa D", "Instalación de Redes", "Finalizado", "2023-08-25"),
            ("5", "Empresa E", "Auditoría de Sistemas", "Trabajo No Iniciado", "2023-07-18"),
            ("6", "Empresa F", "Migración de Datos", "En Proceso", "2023-06-22"),
            ("7", "Empresa G", "Optimización de Procesos", "Completado", "2023-05-30"),
            ("8", "Empresa H", "Soporte Técnico", "Trabajo No Iniciado", "2023-04-12"),
            ("9", "Empresa I", "Desarrollo Web", "En Proceso", "2023-03-28"),
            ("10", "Empresa J", "Evaluación de Seguridad", "Pendiente", "2023-02-11"),
            ("11", "Empresa K", "Diseño Gráfico", "Completado", "2023-01-15"),
            ("12", "Empresa L", "Análisis de Datos", "Finalizado", "2022-12-05"),
            ("13", "Empresa M", "Capacitación de Personal", "Trabajo No Iniciado", "2022-11-22"),
            ("14", "Empresa N", "Optimización de Redes", "En Proceso", "2022-10-13"),
            ("15", "Empresa O", "Desarrollo de App Móvil", "Completado", "2022-09-07"),
            ("16", "Empresa P", "Soporte en Ciberseguridad", "Pendiente", "2022-08-20"),
            ("17", "Empresa Q", "Automatización de Procesos", "Trabajo No Iniciado", "2022-07-11"),
            ("18", "Empresa R", "Asesoría en Transformación Digital", "En Proceso", "2022-06-02"),
            ("19", "Empresa S", "Administración de Sistemas", "Completado", "2022-05-19"),
            ("20", "Empresa T", "Desarrollo de Contenido", "Finalizado", "2022-04-25"),
        ]
        
        muestra_t_orden = datos_t_orden[:50]
        
        for orden in muestra_t_orden:
            t_orden.insert("", "end", values=orden)
            
        scrllbar_t_oc = ttk.Scrollbar(vent_oc, orient="vertical", command=t_orden.yview)
        t_orden.configure(yscrollcommand=scrllbar_t_oc.set)
        scrllbar_t_oc.place(x=1277, y=80, height=479)
            
    def registrar_orden_compra(self):
        
        self.vent_oc.withdraw()
        
        reg_orden = tk.Toplevel(self.vent_oc)
        reg_orden.title("Registrar Orden de Compra")
        reg_orden.geometry("620x380")
        reg_orden.resizable(False, False)
        reg_orden.configure(bg="#373737")
        utils.centrar_ventana(reg_orden)
        reg_orden.protocol("WM_DELETE_WINDOW", lambda: None)
        
        canvas_reg_orden = tk.Canvas(reg_orden, width=620, height=380, bg="#373737", highlightthickness=0)
        canvas_reg_orden.pack()
        
        utils.create_rounded_rectangle(canvas_reg_orden, 10, 10, 610, 194, radius=10, fill="#959595", outline="#959595")
        utils.create_rounded_rectangle(canvas_reg_orden, 10, 204, 610, 330, radius=10, fill="#959595", outline="#959595")
        
        canvas_reg_orden.create_text(20, 20, text="Nro Cotización", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_reg_orden, 20, 38, 170, 68, radius=10, fill="white", outline="#959595")
        reg_cot_or = tk.Entry(reg_orden, font=("Arial", 11), bd=0)
        reg_cot_or.place(x=25, y=43, width=140, height=20)
        
        canvas_reg_orden.create_text(180, 20, text="Nro de Orden de Compra", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_reg_orden, 180, 38, 360, 68, radius=10, fill="white", outline="#959595")
        reg_nr_orden = tk.Entry(reg_orden, font=("Arial", 11), bd=0)
        reg_nr_orden.place(x=185, y=43, width=170, height=20)
        
        canvas_reg_orden.create_text(370, 20, text="Cliente / Empresa", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_reg_orden, 370, 38, 600, 68, radius=10, fill="white", outline="#959595")
        reg_cli_or = tk.Entry(reg_orden, font=("Arial", 11), bd=0)
        reg_cli_or.place(x=375, y=43, width=220, height=20)
        
        canvas_reg_orden.create_text(20, 78, text="Título del Servicio", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_reg_orden, 20, 96, 360, 126, radius=10, fill="white", outline="#959595")
        reg_or_serv = tk.Entry(reg_orden, font=("Arial", 11), bd=0)
        reg_or_serv.place(x=25, y=101, width=330, height=20)
        
        canvas_reg_orden.create_text(370, 78, text="Tiempo de Ejecución", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_reg_orden, 370, 96, 600, 126, radius=10, fill="white", outline="#959595")
        reg_or_tiempo = tk.Entry(reg_orden, font=("Arial", 11), bd=0)
        reg_or_tiempo.place(x=375, y=101, width=220, height=20)
        
        canvas_reg_orden.create_text(20, 136, text="Fecha", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        date_orden = DateEntry(reg_orden, font=("Raleway", 11),state="readonly" , width=17, background='#4A6984', foreground='white', borderwidth=1)
        date_orden.place(x=20, y=154, width=170, height=30)
        
        canvas_reg_orden.create_text(200, 136, text="Estado", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        cbo_est_or = ttk.Combobox(reg_orden, values=["Escoja una Opción", "Trabajo No Iniciado", "En Proceso", "Completado Parcialmente", "Finalizado"], state="readonly", font=("Raleway", 10))
        cbo_est_or.place(x=200, y=154, width=210, height=31)
        cbo_est_or.current(0)
        
        canvas_reg_orden.create_text(20, 214, text="Cotización", anchor="nw", font=("Raleway", 10, "bold"), fill="black")

        btn_ad_or_cot = tk.Button(reg_orden, text="Cambiar documento", command=lambda: utils.adjuntar_archivo(label_or_coti, "cotizacion"), font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_ad_or_cot.place(x=20, y=232, width=140, height=30)
        
        label_or_coti = tk.Label(reg_orden, text="Cotización", font=("Raleway", 9), bg="#373737", fg="white")
        label_or_coti.place(x=170, y=232, width=340, height=30)
        
        btn_ver_or_cot = tk.Button(reg_orden, text="Abrir", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_ver_or_cot.place(x=520, y=232, width=80, height=30)
        
        canvas_reg_orden.create_text(20, 272, text="Orden de Compra", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        
        btn_ad_orden = tk.Button(reg_orden, text="Adjuntar", command=lambda: utils.adjuntar_archivo(label_orden, "orden_compra"), font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_ad_orden.place(x=20, y=290, width=140, height=30)
        
        label_orden = tk.Label(reg_orden, text="Orden de Compra", font=("Raleway", 9), bg="#373737", fg="white")
        label_orden.place(x=170, y=290, width=340, height=30)
        
        btn_ver_orden = tk.Button(reg_orden, text="Abrir", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_ver_orden.place(x=520, y=290, width=80, height=30)
        
        btn_canc_or = tk.Button(reg_orden, text="Cancelar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_canc_or.place(x=205, y=340)
        
        btn_reg_or = tk.Button(reg_orden, text="Registrar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_reg_or.place(x=315, y=340)
        
        utils.aplicar_hover_a_botones([btn_ad_or_cot, btn_ver_or_cot, btn_ad_orden, btn_ver_orden, btn_canc_or, btn_reg_or])
        
    def seguimiento_orden_compra(self):
        
        self.vent_oc.withdraw()
        
        seg_orden = tk.Toplevel(self.vent_oc)
        seg_orden.title("Seguimiento de Orden de Compra")
        seg_orden.geometry("620x438")
        seg_orden.resizable(False, False)
        seg_orden.configure(bg="#373737")
        utils.centrar_ventana(seg_orden)
        seg_orden.protocol("WM_DELETE_WINDOW", lambda: None)
        
        canvas_det_orden = tk.Canvas(seg_orden, width=620, height=438, bg="#373737", highlightthickness=0)
        canvas_det_orden.pack()
        
        utils.create_rounded_rectangle(canvas_det_orden, 10, 10, 610, 194, radius=10, fill="#959595", outline="#959595")
        utils.create_rounded_rectangle(canvas_det_orden, 10, 204, 610, 388, radius=10, fill="#959595", outline="#959595")
        #utils.create_rounded_rectangle(canvas_det_orden, 10, 398, 610, 516, radius=10, fill="#959595", outline="#959595")
        
        canvas_det_orden.create_text(20, 20, text="Nro Cotización", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_det_orden, 20, 38, 170, 68, radius=10, fill="white", outline="#959595")
        det_cot_or = tk.Entry(seg_orden, font=("Arial", 11), bd=0)
        det_cot_or.place(x=25, y=43, width=140, height=20)
        
        canvas_det_orden.create_text(180, 20, text="Nro de Orden de Compra", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_det_orden, 180, 38, 360, 68, radius=10, fill="white", outline="#959595")
        det_nr_orden = tk.Entry(seg_orden, font=("Arial", 11), bd=0)
        det_nr_orden.place(x=185, y=43, width=170, height=20)
        
        canvas_det_orden.create_text(370, 20, text="Cliente / Empresa", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_det_orden, 370, 38, 600, 68, radius=10, fill="white", outline="#959595")
        det_cli_or = tk.Entry(seg_orden, font=("Arial", 11), bd=0)
        det_cli_or.place(x=375, y=43, width=220, height=20)
        
        canvas_det_orden.create_text(20, 78, text="Título del Servicio", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_det_orden, 20, 96, 360, 126, radius=10, fill="white", outline="#959595")
        det_or_serv = tk.Entry(seg_orden, font=("Arial", 11), bd=0)
        det_or_serv.place(x=25, y=101, width=330, height=20)
        
        canvas_det_orden.create_text(370, 78, text="Tiempo de Ejecución", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_det_orden, 370, 96, 600, 126, radius=10, fill="white", outline="#959595")
        det_or_tiempo = tk.Entry(seg_orden, font=("Arial", 11), bd=0)
        det_or_tiempo.place(x=375, y=101, width=220, height=20)
        
        canvas_det_orden.create_text(20, 136, text="Fecha", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        date_det_orden = DateEntry(seg_orden, font=("Raleway", 11),state="readonly" , width=17, background='#4A6984', foreground='white', borderwidth=1)
        date_det_orden.place(x=20, y=154, width=170, height=30)
        
        canvas_det_orden.create_text(200, 136, text="Estado", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        cbo_det_or = ttk.Combobox(seg_orden, values=["Escoja una Opción", "Trabajo No Iniciado", "En Proceso", "Completado Parcialmente", "Finalizado"], state="readonly", font=("Raleway", 10))
        cbo_det_or.place(x=200, y=154, width=210, height=31)
        cbo_det_or.current(0)
        
        canvas_det_orden.create_text(20, 214, text="Cotización", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        
        btn_det_or_cot = tk.Button(seg_orden, text="Cambiar documento", command=lambda: utils.adjuntar_archivo(lbl_det_cot, "cotizacion"), font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_det_or_cot.place(x=20, y=232, width=140, height=30)
        
        lbl_det_cot = tk.Label(seg_orden, text="Cotización", font=("Raleway", 9), bg="#373737", fg="white")
        lbl_det_cot.place(x=170, y=232, width=340, height=30)
        
        btn_ver_cotzz = tk.Button(seg_orden, text="Abrir", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_ver_cotzz.place(x=520, y=232, width=80, height=30)
        
        canvas_det_orden.create_text(20, 272, text="Orden de Compra", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        
        btn_det_orden = tk.Button(seg_orden, text="Cambiar documento", command=lambda: utils.adjuntar_archivo(lbl_det_or, "orden_compra"), font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_det_orden.place(x=20, y=290, width=140, height=30)
        
        lbl_det_or = tk.Label(seg_orden, text="Orden de Compra", font=("Raleway", 9), bg="#373737", fg="white")
        lbl_det_or.place(x=170, y=290, width=340, height=30)
        
        btn_ver_det_or = tk.Button(seg_orden, text="Abrir", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_ver_det_or.place(x=520, y=290, width=80, height=30)
        
        canvas_det_orden.create_text(20, 330, text="Guía de Remisión", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        
        btn_det_gr = tk.Button(seg_orden, text="Adjuntar", command=lambda: utils.adjuntar_archivo(lbl_det_gr, "guia_remision"), font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_det_gr.place(x=20, y=348, width=140, height=30)

        lbl_det_gr = tk.Label(seg_orden, text="Guía de Remisión", font=("Raleway", 9), bg="#373737", fg="white")
        lbl_det_gr.place(x=170, y=348, width=340, height=30)
        
        btn_ver_gr = tk.Button(seg_orden, text="Abrir", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=self.abrir_guias_rem)
        btn_ver_gr.place(x=520, y=348, width=80, height=30)
        
        btn_atras_det = tk.Button(seg_orden, text="Cancelar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_atras_det.place(x=205, y=398)
        
        btn_save_det = tk.Button(seg_orden, text="Guardar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_save_det.place(x=315, y=398)
        
        utils.aplicar_hover_a_botones([btn_det_or_cot, btn_ver_cotzz, btn_det_orden, btn_ver_det_or, btn_det_gr, btn_ver_gr, btn_atras_det, btn_save_det])
    
    # REUTILIZABLE
    def abrir_guias_rem(self):
        vent_guia_rem = tk.Toplevel(self.vent_oc)
        vent_guia_rem.title("")
        vent_guia_rem.geometry("490x256")
        vent_guia_rem.resizable(False, False)
        vent_guia_rem.configure(bg="#373737")
        vent_guia_rem.grab_set()
        utils.centrar_ventana(vent_guia_rem)
        vent_guia_rem.protocol("WM_DELETE_WINDOW", lambda: None)
        
        canvas_det_gr = tk.Canvas(vent_guia_rem, width=490, height=256, bg="#373737", highlightthickness=0)
        canvas_det_gr.pack()
        
        utils.create_rounded_rectangle(canvas_det_gr, 10, 10, 480, 78, radius=10, fill="#959595", outline="#959595")
        
        canvas_det_gr.create_text(20, 20, text="Guía de Remisión", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        btn_elm_gr = tk.Button(vent_guia_rem, text="Eliminar Documento", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_elm_gr.place(x=20, y=38, width=140, height=30)
        
        btn_op_gr = tk.Button(vent_guia_rem, text="Abrir", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_op_gr.place(x=170, y=38, width=80, height=30)
        
        btn_menu_det = tk.Button(vent_guia_rem, text="Cerrar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=vent_guia_rem.destroy)
        btn_menu_det.place(x=195, y=216)
        
        utils.aplicar_hover_a_botones([btn_elm_gr, btn_op_gr, btn_menu_det])
        
        t_guia_rem = ttk.Treeview(vent_guia_rem, columns=("id_gr", "documento"), show="headings", style="Custom.Treeview")
        t_guia_rem.place(x=10, y=88, width=471, height=119)
        
        t_guia_rem.heading("id_gr", text="ID")
        t_guia_rem.heading("documento", text="Documento")
        t_guia_rem.column("id_gr", anchor="center", width=55, stretch=False)
        t_guia_rem.column("documento", anchor="center", width=399, stretch=False)
        
        datos_guia_rem = [
            ("1", "Guía de Despacho A"),
            ("2", "Guía de Despacho B"),
            ("3", "Guía de Despacho C"),
            ("4", "Guía de Despacho D"),
            ("5", "Guía de Despacho E"),
        ]
        
        for dato_orden in datos_guia_rem:
            t_guia_rem.insert("", "end", values=dato_orden)
          
        scrollbar_orden = ttk.Scrollbar(vent_guia_rem, orient="vertical", command=t_guia_rem.yview)
        t_guia_rem.configure(yscrollcommand=scrollbar_orden.set)
        scrollbar_orden.place(x=466, y=88, height=119)


class registrar_factura:
    def __init__(self, root, reg_fact):
        
        self.root = root
        self.reg_fact = reg_fact
        self.root.withdraw()
        
        self.reg_fact = reg_fact
        self.reg_fact.title("Registrar Factura")
        self.reg_fact.geometry("620x496")
        self.reg_fact.resizable(False, False)
        self.reg_fact.configure(bg="#373737")
        utils.centrar_ventana(self.reg_fact)
        self.alerta = alertas(reg_fact)
        self.reg_fact.protocol("WM_DELETE_WINDOW", lambda: None)
        
        canvas_fact = tk.Canvas(reg_fact, width=620, height=496, bg="#373737", highlightthickness=0)
        canvas_fact.pack()
        
        utils.create_rounded_rectangle(canvas_fact, 10, 10, 610, 194, radius=10, fill="#959595", outline="#959595")
        utils.create_rounded_rectangle(canvas_fact, 10, 204, 610, 446, radius=10, fill="#959595", outline="#959595")

        canvas_fact.create_text(20, 20, text="Nro de Orden de Compra", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_fact, 20, 38, 200, 68, radius=10, fill="white", outline="#959595")
        reg_fax_oc = tk.Entry(reg_fact, font=("Arial", 11), bd=0)
        reg_fax_oc.place(x=25, y=43, width=170, height=20)
        
        canvas_fact.create_text(210, 20, text="Nro de Factura", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_fact, 210, 38, 390, 68, radius=10, fill="white", outline="#959595")
        reg_fax = tk.Entry(reg_fact, font=("Arial", 11), bd=0)
        reg_fax.place(x=215, y=43, width=170, height=20)

        canvas_fact.create_text(400, 20, text="Cliente / Empresa", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_fact, 400, 38, 600, 68, radius=10, fill="white", outline="#959595")
        reg_fax_cli = tk.Entry(reg_fact, font=("Arial", 11), bd=0)
        reg_fax_cli.place(x=405, y=43, width=190, height=20)

        canvas_fact.create_text(20, 78, text="Servicio", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_fact, 20, 96, 420, 126, radius=10, fill="white", outline="#959595")
        reg_fax_serv = tk.Entry(reg_fact, font=("Arial", 11), bd=0)
        reg_fax_serv.place(x=25, y=101, width=390, height=20)

        canvas_fact.create_text(430, 78, text="Fecha de Emisión", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        fecha_em_reg_fact = DateEntry(reg_fact, font=("Raleway", 11),state="readonly" , width=17, background='#4A6984', foreground='white', borderwidth=1)
        fecha_em_reg_fact.place(x=430, y=96, width=170, height=30)
        
        canvas_fact.create_text(20, 136, text="Fecha de Vencimiento", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        fecha_ven_reg_fact = DateEntry(reg_fact, font=("Raleway", 11),state="readonly" , width=17, background='#4A6984', foreground='white', borderwidth=1)
        fecha_ven_reg_fact.place(x=20, y=154, width=170, height=30)

        canvas_fact.create_text(200, 136, text="Forma de Pago", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        cbo_fp_fact = ttk.Combobox(reg_fact, values=["Escoja una Opción", "30 días", "60 días", "90 días", "180 días", "Al Contado"], state="readonly", font=("Raleway", 10))
        cbo_fp_fact.place(x=200, y=154, width=135, height=31)
        cbo_fp_fact.current(0)
        
        canvas_fact.create_text(345, 136, text="Estado", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        cbo_est_fact = ttk.Combobox(reg_fact, values=["Escoja una Opción", "Cancelado", "Pendiente", "Falta Pagar"], state="readonly", font=("Raleway", 10))
        cbo_est_fact.place(x=345, y=154, width=135, height=31)
        cbo_est_fact.current(0)

        canvas_fact.create_text(20, 214, text="Informe Técnico", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        btn_ad_it = tk.Button(reg_fact, text="Adjuntar", command=lambda: utils.adjuntar_archivo(label_it, "informe_tecnico"), font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_ad_it.place(x=20, y=232, width=140, height=30)
        label_it = tk.Label(reg_fact, text="Informe Técnico", font=("Raleway", 9), bg="#373737", fg="white")
        label_it.place(x=170, y=232, width=340, height=30)
        btn_op_it = tk.Button(reg_fact, text="Abrir", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=self.abrir_informe_tecnico)
        btn_op_it.place(x=520, y=232, width=80, height=30)
        
        canvas_fact.create_text(20, 272, text="Planos", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        btn_ad_pl = tk.Button(reg_fact, text="Adjuntar", command=lambda: utils.adjuntar_archivo(label_pl, "doc_planos"), font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_ad_pl.place(x=20, y=290, width=140, height=30)
        label_pl = tk.Label(reg_fact, text="Planos", font=("Raleway", 9), bg="#373737", fg="white")
        label_pl.place(x=170, y=290, width=340, height=30)
        btn_op_pl = tk.Button(reg_fact, text="Abrir", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=self.abrir_planos)
        btn_op_pl.place(x=520, y=290, width=80, height=30)
        
        canvas_fact.create_text(20, 330, text="Acta de Conformidad", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        btn_ad_ac = tk.Button(reg_fact, text="Adjuntar", command=lambda: utils.adjuntar_archivo(label_ac, "acta_conformidad"), font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_ad_ac.place(x=20, y=348, width=140, height=30)
        label_ac = tk.Label(reg_fact, text="Acta de Conformidad", font=("Raleway", 9), bg="#373737", fg="white")
        label_ac.place(x=170, y=348, width=340, height=30)
        btn_op_ac = tk.Button(reg_fact, text="Abrir", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=self.abrir_acta_conformidad)
        btn_op_ac.place(x=520, y=348, width=80, height=30)
        
        canvas_fact.create_text(20, 388, text="Factura", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        btn_ad_fact = tk.Button(reg_fact, text="Adjuntar", command=lambda: utils.adjuntar_archivo(label_fact, "doc_factura"), font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_ad_fact.place(x=20, y=406, width=140, height=30)
        label_fact = tk.Label(reg_fact, text="Factura", font=("Raleway", 9), bg="#373737", fg="white")
        label_fact.place(x=170, y=406, width=340, height=30)
        btn_op_fact = tk.Button(reg_fact, text="Abrir", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_op_fact.place(x=520, y=406, width=80, height=30)

        btn_can_fact = tk.Button(reg_fact, text="Cancelar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_can_fact.place(x=205, y=456)
        
        btn_save_regfact = tk.Button(reg_fact, text="Registrar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_save_regfact.place(x=315, y=456)
        
        utils.aplicar_hover_a_botones([btn_ad_it, btn_op_it, btn_ad_pl, btn_op_pl, btn_ad_ac, btn_op_ac, btn_ad_fact, btn_op_fact, btn_can_fact, btn_save_regfact])

    # REUTILIZABLE
    def abrir_informe_tecnico(self):
        vent_informe = tk.Toplevel(self.reg_fact)
        vent_informe.title
        vent_informe.geometry("490x256")
        vent_informe.resizable(False, False)
        vent_informe.configure(bg="#373737")
        vent_informe.grab_set()
        utils.centrar_ventana(vent_informe)
        vent_informe.protocol("WM_DELETE_WINDOW", lambda: None)

        canvas_inf = tk.Canvas(vent_informe, width=490, height=256, bg="#373737", highlightthickness=0)
        canvas_inf.pack()

        utils.create_rounded_rectangle(canvas_inf, 10, 10, 480, 78, radius=10, fill="#959595", outline="#959595")

        canvas_inf.create_text(20, 20, text="Informe Técnico", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        btn_elm_if = tk.Button(vent_informe, text="Eliminar Documento", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_elm_if.place(x=20, y=38, width=140, height=30)
        
        btn_abrir_if = tk.Button(vent_informe, text="Abrir", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_abrir_if.place(x=170, y=38, width=80, height=30)
        
        btn_menu_if = tk.Button(vent_informe, text="Cerrar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=vent_informe.destroy)
        btn_menu_if.place(x=195, y=216)
        
        utils.aplicar_hover_a_botones([btn_elm_if, btn_abrir_if, btn_menu_if])

        t_informe = ttk.Treeview(vent_informe, columns=("id_it", "doc_if"), show="headings", style="Custom.Treeview")
        t_informe.place(x=10, y=88, width=471, height=119)

        t_informe.heading("id_it", text="ID")
        t_informe.heading("doc_if", text="Documento")
        t_informe.column("id_it", anchor="center", width=55, stretch=False)
        t_informe.column("doc_if", anchor="center", width=399, stretch=False)

        datos_informe = [
            ("1", "Informe Técnico A"),
            ("2", "Informe Técnico B"),
            ("3", "Informe Técnico C"),
            ("4", "Informe Técnico D"),
            ("5", "Informe Técnico E"),
        ]

        for dato_orden_if in datos_informe:
            t_informe.insert("", "end", values=dato_orden_if)

        scrollbar_informe = ttk.Scrollbar(vent_informe, orient="vertical", command=t_informe.yview)
        t_informe.configure(yscrollcommand=scrollbar_informe.set)
        scrollbar_informe.place(x=466, y=88, height=119)

    # REUTILIZABLE
    def abrir_planos(self):
        vent_planos = tk.Toplevel(self.reg_fact)
        vent_planos.title
        vent_planos.geometry("490x256")
        vent_planos.resizable(False, False)
        vent_planos.configure(bg="#373737")
        vent_planos.grab_set()
        utils.centrar_ventana(vent_planos)
        vent_planos.protocol("WM_DELETE_WINDOW", lambda: None)

        canvas_plan = tk.Canvas(vent_planos, width=490, height=256, bg="#373737", highlightthickness=0)
        canvas_plan.pack()

        utils.create_rounded_rectangle(canvas_plan, 10, 10, 480, 78, radius=10, fill="#959595", outline="#959595")

        canvas_plan.create_text(20, 20, text="Planos", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        btn_elm_pl = tk.Button(vent_planos, text="Eliminar Documento", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_elm_pl.place(x=20, y=38, width=140, height=30)
        
        btn_abrir_pl = tk.Button(vent_planos, text="Abrir", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_abrir_pl.place(x=170, y=38, width=80, height=30)
        
        btn_menu_pl = tk.Button(vent_planos, text="Cerrar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=vent_planos.destroy)
        btn_menu_pl.place(x=195, y=216)
        
        utils.aplicar_hover_a_botones([btn_elm_pl, btn_abrir_pl, btn_menu_pl])

        t_planos = ttk.Treeview(vent_planos, columns=("id_pl", "doc_pl"), show="headings", style="Custom.Treeview")
        t_planos.place(x=10, y=88, width=471, height=119)

        t_planos.heading("id_pl", text="ID")
        t_planos.heading("doc_pl", text="Documento")
        t_planos.column("id_pl", anchor="center", width=55, stretch=False)
        t_planos.column("doc_pl", anchor="center", width=399, stretch=False)

        datos_planos = [
            ("1", "Planos A"),
            ("2", "Planos B"),
            ("3", "Planos C"),
            ("4", "Planos D"),
            ("5", "Planos E"),
        ]

        for dato_orden_pl in datos_planos:
            t_planos.insert("", "end", values=dato_orden_pl)

        scrollbar_planos = ttk.Scrollbar(vent_planos, orient="vertical", command=t_planos.yview)
        t_planos.configure(yscrollcommand=scrollbar_planos.set)
        scrollbar_planos.place(x=466, y=88, height=119)

    # REUTILIZABLE
    def abrir_acta_conformidad(self):
        vent_acta = tk.Toplevel(self.reg_fact)
        vent_acta.title
        vent_acta.geometry("490x256")
        vent_acta.resizable(False, False)
        vent_acta.configure(bg="#373737")
        vent_acta.grab_set()
        utils.centrar_ventana(vent_acta)
        vent_acta.protocol("WM_DELETE_WINDOW", lambda: None)

        canvas_acta = tk.Canvas(vent_acta, width=490, height=256, bg="#373737", highlightthickness=0)
        canvas_acta.pack()

        utils.create_rounded_rectangle(canvas_acta, 10, 10, 480, 78, radius=10, fill="#959595", outline="#959595")

        canvas_acta.create_text(20, 20, text="Acta de Conformidad", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        btn_elm_ac = tk.Button(vent_acta, text="Eliminar Documento", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_elm_ac.place(x=20, y=38, width=140, height=30)
        
        btn_abrir_ac = tk.Button(vent_acta, text="Abrir", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_abrir_ac.place(x=170, y=38, width=80, height=30)
        
        btn_menu_ac = tk.Button(vent_acta, text="Cerrar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=vent_acta.destroy)
        btn_menu_ac.place(x=195, y=216)
        
        utils.aplicar_hover_a_botones([btn_elm_ac, btn_abrir_ac, btn_menu_ac]) 

        t_acta = ttk.Treeview(vent_acta, columns=("id_ac", "doc_ac"), show="headings", style="Custom.Treeview")
        t_acta.place(x=10, y=88, width=471, height=119)

        t_acta.heading("id_ac", text="ID")
        t_acta.heading("doc_ac", text="Documento")
        t_acta.column("id_ac", anchor="center", width=55, stretch=False)
        t_acta.column("doc_ac", anchor="center", width=399, stretch=False)

        datos_acta = [
            ("1", "Acta de Conformidad A"),
            ("2", "Acta de Conformidad B"),
            ("3", "Acta de Conformidad C"),
            ("4", "Acta de Conformidad D"),
            ("5", "Acta de Conformidad E"),
        ]

        for dato_orden_ac in datos_acta:
            t_acta.insert("", "end", values=dato_orden_ac)

        scrollbar_acta = ttk.Scrollbar(vent_acta, orient="vertical", command=t_acta.yview)
        t_acta.configure(yscrollcommand=scrollbar_acta.set)
        scrollbar_acta.place(x=466, y=88, height=119)


class detalles_factura:
    def __init__(self, root, seg_fact):
        
        self.root = root
        self.seg_fact = seg_fact
        self.root.withdraw()
        
        self.seg_fact = seg_fact
        self.seg_fact.title("Seguimiento de Facturas")
        self.seg_fact.geometry("620x670")
        self.seg_fact.resizable(False, False)
        self.seg_fact.configure(bg="#373737")
        utils.centrar_ventana(self.seg_fact)
        self.alerta = alertas(seg_fact)
        self.seg_fact.protocol("WM_DELETE_WINDOW", lambda: None)
    
        canvas_seg_fax = tk.Canvas(seg_fact, width=620, height=670, bg="#373737", highlightthickness=0)
        canvas_seg_fax.pack()
        
        utils.create_rounded_rectangle(canvas_seg_fax, 10, 10, 610, 194, radius=10, fill="#959595", outline="#959595")
        utils.create_rounded_rectangle(canvas_seg_fax, 10, 204, 610, 620, radius=10, fill="#959595", outline="#959595")
        
        canvas_seg_fax.create_text(20, 20, text="Nro de Orden de Compra", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_seg_fax, 20, 38, 200, 68, radius=10, fill="white", outline="#959595")
        seg_fax_oc = tk.Entry(seg_fact, font=("Arial", 11), bd=0)
        seg_fax_oc.place(x=25, y=43, width=170, height=20)
        
        canvas_seg_fax.create_text(210, 20, text="Nro de Factura", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_seg_fax, 210, 38, 390, 68, radius=10, fill="white", outline="#959595")
        seg_fax = tk.Entry(seg_fact, font=("Arial", 11), bd=0)
        seg_fax.place(x=215, y=43, width=170, height=20)
        
        canvas_seg_fax.create_text(400, 20, text="Cliente / Empresa", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_seg_fax, 400, 38, 600, 68, radius=10, fill="white", outline="#959595")
        seg_fax_cli = tk.Entry(seg_fact, font=("Arial", 11), bd=0)
        seg_fax_cli.place(x=405, y=43, width=190, height=20)
        
        canvas_seg_fax.create_text(20, 78, text="Servicio", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        utils.create_rounded_rectangle(canvas_seg_fax, 20, 96, 420, 126, radius=10, fill="white", outline="#959595")
        seg_fax_serv = tk.Entry(seg_fact, font=("Arial", 11), bd=0)
        seg_fax_serv.place(x=25, y=101, width=390, height=20)
        
        canvas_seg_fax.create_text(430, 78, text="Fecha de Emisión", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        fecha_em_det_fact = DateEntry(seg_fact, font=("Raleway", 11),state="readonly" , width=17, background='#4A6984', foreground='white', borderwidth=1)
        fecha_em_det_fact.place(x=430, y=96, width=170, height=30)
        
        canvas_seg_fax.create_text(20, 136, text="Fecha de Vencimiento", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        fecha_ven_det_fact = DateEntry(seg_fact, font=("Raleway", 11),state="readonly" , width=17, background='#4A6984', foreground='white', borderwidth=1)
        fecha_ven_det_fact.place(x=20, y=154, width=170, height=30)
        
        canvas_seg_fax.create_text(200, 136, text="Forma de Pago", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        cbo_fp_seg_fact = ttk.Combobox(seg_fact, values=["Escoja una Opción", "30 días", "60 días", "90 días", "180 días", "Al Contado"], state="readonly", font=("Raleway", 10))
        cbo_fp_seg_fact.place(x=200, y=154, width=135, height=31)
        cbo_fp_seg_fact.current(0)
        
        canvas_seg_fax.create_text(345, 136, text="Estado", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        cbo_est_seg_fact = ttk.Combobox(seg_fact, values=["Escoja una Opción", "Cancelado", "Pendiente", "Falta Pagar"], state="readonly", font=("Raleway", 10))
        cbo_est_seg_fact.place(x=345, y=154, width=135, height=31)
        cbo_est_seg_fact.current(0)
        
        canvas_seg_fax.create_text(20, 214, text="Cotización", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        btn_ch_coti = tk.Button(seg_fact, text="Cambiar documento", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=lambda: utils.adjuntar_archivo(lbl_ch_coti, "cotizacion"))
        btn_ch_coti.place(x=20, y=232, width=140, height=30)
        lbl_ch_coti = tk.Label(seg_fact, text="Cotización", font=("Raleway", 9), bg="#373737", fg="white")
        lbl_ch_coti.place(x=170, y=232, width=340, height=30)
        btn_op_doc_cot = tk.Button(seg_fact, text="Abrir", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_op_doc_cot.place(x=520, y=232, width=80, height=30)
        
        canvas_seg_fax.create_text(20, 272, text="Orden de Compra", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        btn_ch_oc = tk.Button(seg_fact, text="Cambiar documento", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=lambda: utils.adjuntar_archivo(lbl_ch_oc, "orden_compra"))
        btn_ch_oc.place(x=20, y=290, width=140, height=30)
        lbl_ch_oc = tk.Label(seg_fact, text="Orden de Compra", font=("Raleway", 9), bg="#373737", fg="white")
        lbl_ch_oc.place(x=170, y=290, width=340, height=30)
        btn_op_doc_oc = tk.Button(seg_fact, text="Abrir", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_op_doc_oc.place(x=520, y=290, width=80, height=30)
        
        canvas_seg_fax.create_text(20, 330, text="Guía de Remisión", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        btn_ch_gr = tk.Button(seg_fact, text="Cambiar documento", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=lambda: utils.adjuntar_archivo(lbl_ch_gr, "guia_remision"))
        btn_ch_gr.place(x=20, y=348, width=140, height=30)
        lbl_ch_gr = tk.Label(seg_fact, text="Guía de Remisión", font=("Raleway", 9), bg="#373737", fg="white")
        lbl_ch_gr.place(x=170, y=348, width=340, height=30)
        btn_op_doc_gr = tk.Button(seg_fact, text="Abrir", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_op_doc_gr.place(x=520, y=348, width=80, height=30)
        
        canvas_seg_fax.create_text(20, 388, text="Informe Técnico", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        btn_ch_it = tk.Button(seg_fact, text="Cambiar documento", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=lambda: utils.adjuntar_archivo(lbl_ch_it, "informe_tecnico"))
        btn_ch_it.place(x=20, y=406, width=140, height=30)
        lbl_ch_it = tk.Label(seg_fact, text="Informe Técnico", font=("Raleway", 9), bg="#373737", fg="white")
        lbl_ch_it.place(x=170, y=406, width=340, height=30)
        btn_op_doc_it = tk.Button(seg_fact, text="Abrir", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_op_doc_it.place(x=520, y=406, width=80, height=30)
        
        canvas_seg_fax.create_text(20, 446, text="Planos", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        btn_ch_pl = tk.Button(seg_fact, text="Cambiar documento", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=lambda: utils.adjuntar_archivo(lbl_ch_pl, "doc_planos"))
        btn_ch_pl.place(x=20, y=464, width=140, height=30)
        lbl_ch_pl = tk.Label(seg_fact, text="Planos", font=("Raleway", 9), bg="#373737", fg="white")
        lbl_ch_pl.place(x=170, y=464, width=340, height=30)
        btn_op_doc_pl = tk.Button(seg_fact, text="Abrir", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_op_doc_pl.place(x=520, y=464, width=80, height=30)
        
        canvas_seg_fax.create_text(20, 504, text="Acta de Conformidad", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        btn_ch_ac = tk.Button(seg_fact, text="Cambiar documento", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=lambda: utils.adjuntar_archivo(lbl_ch_ac, "acta_conformidad"))
        btn_ch_ac.place(x=20, y=522, width=140, height=30)
        lbl_ch_ac = tk.Label(seg_fact, text="Acta de Conformidad", font=("Raleway", 9), bg="#373737", fg="white")
        lbl_ch_ac.place(x=170, y=522, width=340, height=30)
        btn_op_doc_ac = tk.Button(seg_fact, text="Abrir", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_op_doc_ac.place(x=520, y=522, width=80, height=30)
        
        canvas_seg_fax.create_text(20, 562, text="Factura", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        btn_add_fact = tk.Button(seg_fact, text="Adjuntar", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=lambda: utils.adjuntar_archivo(lbl_add_fact, "doc_factura"))
        btn_add_fact.place(x=20, y=580, width=140, height=30)
        lbl_add_fact = tk.Label(seg_fact, text="Factura", font=("Raleway", 9), bg="#373737", fg="white")
        lbl_add_fact.place(x=170, y=580, width=340, height=30)
        btn_op_doc_fac = tk.Button(seg_fact, text="Abrir", font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_op_doc_fac.place(x=520, y=580, width=80, height=30)
        
        btn_can_segf = tk.Button(seg_fact, text="Cancelar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_can_segf.place(x=205, y=630)
        
        btn_save_segf = tk.Button(seg_fact, text="Guardar", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_save_segf.place(x=315, y=630)
        
        utils.aplicar_hover_a_botones([btn_ch_coti, btn_op_doc_cot, btn_ch_oc, btn_op_doc_oc, 
                                       btn_ch_gr, btn_op_doc_gr, btn_ch_it, btn_op_doc_it, 
                                       btn_ch_pl, btn_op_doc_pl, btn_ch_ac, btn_op_doc_ac, 
                                       btn_add_fact, btn_op_doc_fac, btn_can_segf, btn_save_segf])


class buscador:
    def __init__(self, root, buscar_doc):
        
        self.root = root
        self.buscar_doc = buscar_doc
        self.root.withdraw()
        
        self.buscar_doc = buscar_doc
        self.buscar_doc.title("Buscar Documentos")
        self.buscar_doc.geometry("1100x664")
        self.buscar_doc.resizable(False, False)
        self.buscar_doc.configure(bg="#373737")
        utils.centrar_ventana(self.buscar_doc)
        self.alerta = alertas(buscar_doc)
        
        self.buscar_doc.protocol("WM_DELETE_WINDOW", lambda: None)
        
        canvas_buscar_doc = tk.Canvas(buscar_doc, width=1100, height=664, bg="#373737", highlightthickness=0)
        canvas_buscar_doc.pack()
        
        utils.create_rounded_rectangle(canvas_buscar_doc, 10, 10, 300, 126, radius=10, fill="#959595", outline="#959595")
        utils.create_rounded_rectangle(canvas_buscar_doc, 310, 66, 1090, 126, radius=10, fill="#959595", outline="#959595")
        #utils.create_rounded_rectangle(canvas_buscar_doc, 10, 136, 1090, 614, radius=10, fill="#959595", outline="#959595")

        canvas_buscar_doc.create_text(523, 15, text="BUSCAR DOCUMENTOS", anchor="nw", font=("Raleway", 20, "bold"), fill="White")
        canvas_buscar_doc.create_text(20, 20, text="Filtro", anchor="nw", font=("Raleway", 20, "bold"), fill="White")
        
        canvas_buscar_doc.create_text(20, 68, text="Por Tipo de Documentos", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        cbo_tipo_doc = ttk.Combobox(buscar_doc, values=["Todos los registros", "Cotización", "Orden de Compra", "Guía de Remisión", "Planos", "Acta de Conformidad", "Informe Técnico", "Factura"], state="readonly", font=("Arial", 10))
        cbo_tipo_doc.place(x=20, y=86, width=270, height=30)
        cbo_tipo_doc.current(0)
        
        search_canvas_buscador = tk.Canvas(buscar_doc, width=760, height=40, bg="#373737", highlightthickness=0)
        search_canvas_buscador.place(x=320, y=76)
        
        utils.create_rounded_rectangle(search_canvas_buscador, 0, 0, 760, 40, radius=10, fill="white", outline="#959595")
        search_canvas_buscador.create_line(725, 7, 725, 34, fill="gray", width=2)
        
        search_icon_path = os.path.join(ICON_DIR, "search.png")
        try:
            search_icon_buscador = tk.PhotoImage(file=search_icon_path)
            search_icon_id_buscador = search_canvas_buscador.create_image(731, 8, anchor="nw", image=search_icon_buscador)
            search_canvas_buscador.image = search_icon_buscador
        except Exception as e:
            raise FileNotFoundError(f"El archivo del icono no se encontró en la ruta: {search_icon_path}. Error: {e}")

        search_canvas_buscador.tag_bind(search_icon_id_buscador, "<Button-1>", lambda e: self.alerta.no_datos())
        search_entry_buscador = tk.Entry(search_canvas_buscador, font=("Arial", 13), width=40, bd=0, relief="flat", fg='grey')
        search_entry_buscador.insert(0, "Buscar...")
        search_entry_buscador.bind("<FocusIn>", lambda event: utils.clear_placeholder(event, search_entry_buscador))
        search_entry_buscador.bind("<FocusOut>", lambda event: utils.placeholder_search(event, search_entry_buscador))
        search_entry_buscador.place(x=6, y=7, width=711, height=27)

        t_buscadoc = ttk.Treeview(buscar_doc, columns=("id_doc", "doc", "t_doc", "date_doc", "ruta"), show="headings", style="Custom.Treeview")
        t_buscadoc.place(x=10, y=136, width=1081, height=479)
        
        t_buscadoc.heading("id_doc", text="ID")
        t_buscadoc.heading("doc", text="Documento")
        t_buscadoc.heading("t_doc", text="Tipo de Documento")
        t_buscadoc.heading("date_doc", text="Fecha")
        t_buscadoc.heading("ruta", text="Ruta")
        
        t_buscadoc.column("id_doc", anchor="center", width=75, stretch=False)
        t_buscadoc.column("doc", anchor="center", width=302, stretch=False)
        t_buscadoc.column("t_doc", anchor="center", width=140, stretch=False)
        t_buscadoc.column("date_doc", anchor="center", width=100, stretch=False)
        t_buscadoc.column("ruta", anchor="center", width=448, stretch=False)
        
        datos_t_buscadoc = [
            ("1", "Factura_001", "Factura", "2023-11-01", "C:/Documentos/Facturas/Factura_001.pdf"),
            ("2", "Contrato_2023", "Contrato", "2023-10-15", "C:/Documentos/Contratos/Contrato_2023.pdf"),
            ("3", "Informe_TI", "Informe Técnico", "2023-09-20", "D:/Proyectos/Informes/Informe_TI.docx"),
            ("4", "Plan_Marketing", "Plan", "2023-08-25", "E:/Marketing/Planes/Plan_Marketing.xlsx"),
            ("5", "Guía_Despacho", "Guía de Remisión", "2023-07-18", "C:/Despachos/Guias/Guía_Despacho.pdf"),
            ("6", "Acta_Reunión", "Acta de Reunión", "2023-06-22", "C:/Reuniones/Actas/Acta_Reunión.docx"),
            ("7", "Propuesta_Cliente", "Propuesta", "2023-05-30", "D:/Ventas/Propuestas/Propuesta_Cliente.pdf"),
            ("8", "Manual_Usuario", "Manual", "2023-04-12", "E:/Soporte/Manuales/Manual_Usuario.pdf"),
            ("9", "Reporte_Anual", "Reporte", "2023-03-28", "C:/Reportes/Anuales/Reporte_Anual.pdf"),
            ("10", "Evaluación_Técnica", "Evaluación", "2023-02-11", "D:/Evaluaciones/Técnicas/Evaluación_Técnica.pdf"),
            ("11", "Lista_Precios", "Lista de Precios", "2023-01-15", "C:/Precios/Listas/Lista_Precios.xlsx"),
            ("12", "Política_Calidad", "Política", "2022-12-05", "D:/Calidad/Políticas/Política_Calidad.pdf"),
            ("13", "Memorandum_Interno", "Memorandum", "2022-11-22", "C:/Memorandos/Memorandum_Interno.docx"),
            ("14", "Certificado_Garantía", "Certificado", "2022-10-13", "E:/Garantías/Certificados/Certificado_Garantía.pdf"),
            ("15", "Acta_Cierre", "Acta de Cierre", "2022-09-07", "C:/Proyectos/Cierre/Acta_Cierre.pdf"),
            ("16", "Solicitud_Cliente", "Solicitud", "2022-08-20", "D:/Clientes/Solicitudes/Solicitud_Cliente.docx"),
            ("17", "Protocolo_Seguridad", "Protocolo", "2022-07-11", "C:/Seguridad/Protocolos/Protocolo_Seguridad.pdf"),
            ("18", "Guía_Operativa", "Guía", "2022-06-02", "E:/Operaciones/Guías/Guía_Operativa.pdf"),
            ("19", "Plan_Financiero", "Plan", "2022-05-19", "C:/Finanzas/Planes/Plan_Financiero.xlsx"),
            ("20", "Hoja_Cotización", "Cotización", "2022-04-25", "D:/Cotizaciones/Hojas/Hoja_Cotización.pdf"),
            ("21", "Informe_Progreso", "Informe", "2022-03-15", "C:/Informes/Progreso/Informe_Progreso.docx"),
            ("22", "Acta_Conformidad", "Acta de Conformidad", "2022-02-10", "C:/Actas/Conformidad/Acta_Conformidad.pdf"),
            ("23", "Plan_Estrategico", "Plan", "2022-01-05", "D:/Estrategia/Planes/Plan_Estrategico.xlsx"),
            ("24", "Certificado_Trabajo", "Certificado", "2021-12-20", "C:/Certificados/Trabajo/Certificado_Trabajo.pdf"),
            ("25", "Informe_Financiero", "Informe", "2021-11-15", "D:/Finanzas/Informes/Informe_Financiero.docx"),
        ]
        
        datos_mostrados = datos_t_buscadoc[:100]
        
        for dato in datos_mostrados:
            t_buscadoc.insert("", "end", values=dato)
        
        scrllbar_t_busq = ttk.Scrollbar(buscar_doc, orient="vertical", command=t_buscadoc.yview)
        t_buscadoc.configure(yscrollcommand=scrllbar_t_busq.set)
        scrllbar_t_busq.place(x=1077, y=136, height=479)
        
        cbo_busq_page = ttk.Combobox(buscar_doc, values=["1", "2"], state="readonly", font=("Arial", 10))
        cbo_busq_page.place(x=10, y=624, width=70, height=30)
        cbo_busq_page.current(0)
        
        btn_menu_busq = tk.Button(buscar_doc, text="Atrás", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_menu_busq.place(x=500, y=624)
        
        utils.aplicar_hover_a_botones([btn_menu_busq])


class ventana_inicio:
    def __init__(self, root):
        self.root = root
        self.root.title("Inicio")
        self.root.geometry("1400x700")
        self.root.resizable(False, False)
        self.root.configure(bg="#373737")
        utils.centrar_ventana(root)
        
        self.alerta = alertas(root)
        
        # Usar el método cerrar_prog de alertas para el cierre de la ventana principal
        self.root.protocol("WM_DELETE_WINDOW", self.alerta.cerrar_prog)

        canvas = tk.Canvas(root, width=1400, height=700, bg="#373737", highlightthickness=0)
        canvas.pack()

        utils.create_rounded_rectangle(canvas, 10, 10, 300, 558, radius=10, fill="#959595", outline="#959595")
        utils.create_rounded_rectangle(canvas, 10, 568, 300, 690, radius=10, fill="#959595", outline="#959595")
        #utils.create_rounded_rectangle(canvas, 310, 80, 1390, 650, radius=10, fill="#959595", outline="#959595")

        canvas.create_text(20, 22, text="Opciones", anchor="nw", font=("Raleway", 20, "bold"), fill="White")

        # Botones opcion
        btn_gen_cot = tk.Button(root, text="Generar Cotización", width=37, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=self.abrir_ventana_cotizacion)
        btn_gen_cot.place(x=22, y=80)

        btn_reg_cli = tk.Button(root, text="Ver Registro de Clientes", width=37, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=self.abrir_ventana_clientes)
        btn_reg_cli.place(x=22, y=125)

        btn_reg_cot = tk.Button(root, text="Ver Registro de Cotización", width=37, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=self.abrir_reg_cotizacion)
        btn_reg_cot.place(x=22, y=170)

        btn_reg_oc = tk.Button(root, text="Ver Registro de Orden de Compra", width=37, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=self.abrir_orden_compra)
        btn_reg_oc.place(x=22, y=215)

        btn_reg_fact = tk.Button(root, text="Registrar Factura", width=37, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=self.abrir_registrar_factura)
        btn_reg_fact.place(x=22, y=260)

        btn_seg_fact = tk.Button(root, text="Seguimiento de Factura", width=37, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=self.abrir_detalles_factura)
        btn_seg_fact.place(x=22, y=305)

        btn_search = tk.Button(root, text="Buscar Documentos", width=37, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white", command=self.abrir_buscador)
        btn_search.place(x=22, y=350)

        btn_act = tk.Button(root, text="Actualizar Registro", width=37, height=1, font=("Raleway", 9), command=self.alerta.confirmacion_cotizacion, activebackground="#7F7F7F", activeforeground="white")
        btn_act.place(x=22, y=395)
        
        btn_op_file_fact = tk.Button(self.root, text="Abrir Carpeta", width=37, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_op_file_fact.place(x=22, y=440)

        btn_exit = tk.Button(self.root, text="Salir", width=37, height=1, font=("Raleway", 9), command=self.alerta.cerrar_prog, activebackground="#7F7F7F", activeforeground="white")
        btn_exit.place(x=22, y=485)
        
        btn_siguiente = tk.Button(root, text="Siguiente", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_siguiente.place(x=1290, y=660)
        
        btn_atras = tk.Button(root, text="Anterior", width=13, height=1, font=("Raleway", 9), activebackground="#7F7F7F", activeforeground="white")
        btn_atras.place(x=1180, y=660)

        # Vincular los eventos a cada botón
        utils.aplicar_hover_a_botones([btn_gen_cot, btn_reg_cli, btn_reg_cot, btn_reg_oc, btn_reg_fact, btn_seg_fact, btn_search, btn_act, btn_op_file_fact, btn_exit, btn_siguiente, btn_atras])

        canvas.create_text(20, 578, text="Filtro", anchor="nw", font=("Raleway", 20, "bold"), fill="White")

        canvas.create_text(20, 632, text="Por Estado", anchor="nw", font=("Raleway", 10, "bold"), fill="black")
        
        cbo_fact = ttk.Combobox(root, values=["Todos los registros", "Cancelado", "Pendiente", "Falta Pagar"], state="readonly", font=("Raleway", 10))
        cbo_fact.place(x=20, y=650, width=270, height=30)
        cbo_fact.current(0)
        
        cbo_page = ttk.Combobox(root, values=["1", "2"], state="readonly", font=("Arial", 10))
        cbo_page.place(x=310, y=660, width=70, height=30)
        cbo_page.current(0)

        canvas.create_text(400, 23, text="SOLUCIONES PLÁSTICAS METÁLICAS SAC", anchor="nw", font=("Raleway", 21, "bold"), fill="White")
        
        search_canvas = tk.Canvas(root, width=350, height=40, bg="#373737", highlightthickness=0)
        search_canvas.place(x=1040, y=20)
        
        utils.create_rounded_rectangle(search_canvas, 0, 0, 350, 40, radius=10, fill="white", outline="gray")
        search_canvas.create_line(315, 7, 315, 34, fill="gray", width=2)
        
        search_icon_path = os.path.join(ICON_DIR, "search.png")  # Ruta dinámica
        try:
            search_icon = tk.PhotoImage(file=search_icon_path)
            search_icon_id = search_canvas.create_image(321, 8, anchor="nw", image=search_icon)
            search_canvas.image = search_icon
        except Exception as e:
            raise FileNotFoundError(f"El archivo del icono no se encontró en la ruta: {search_icon_path}. Error: {e}")

        search_canvas.tag_bind(search_icon_id, "<Button-1>", lambda e: self.alerta.no_datos())
        
        search_entry = tk.Entry(search_canvas, font=("Arial", 13), width=40, bd=0, relief="flat", fg='grey')
        search_entry.insert(0, "Buscar...")
        search_entry.bind("<FocusIn>", lambda event: utils.clear_placeholder(event, search_entry))  # Limpiar cuando el usuario hace clic
        search_entry.bind("<FocusOut>", lambda event: utils.placeholder_search(event, search_entry))  # Volver a mostrar si está vacío
        search_entry.place(x=6, y=7, width=301, height=27)
        
        style = ttk.Style()
        style.theme_use("clam")
        
        style.configure("Custom.Treeview.Heading", background="lightblue", foreground="black", font=("Arial", 9))
        style.configure("Custom.Treeview", background="white", foreground="black", rowheight=30)

        tree = ttk.Treeview(root, columns=("oc", "fact", "cliente", "descripcion", "estado_fac", "fecha_em", "fecha_ven"), show="headings", style="Custom.Treeview")
        tree.place(x=310, y=80, width=1081, height=569)

        tree.heading("oc", text="Orden de Compra")
        tree.heading("fact", text="Factura")
        tree.heading("cliente", text="Cliente / Empresa")
        tree.heading("descripcion", text="Descripción / Servicio")
        tree.heading("estado_fac", text="Estado")
        tree.heading("fecha_em", text="Fecha de Em.")
        tree.heading("fecha_ven", text="Fecha de Venc.")

        tree.column("oc", anchor="center", width=115, stretch=False)
        tree.column("fact", anchor="center", width=110, stretch=False)
        tree.column("cliente", anchor="center", width=168, stretch=False)
        tree.column("descripcion", anchor="center", width=352, stretch=False)
        tree.column("estado_fac", anchor="center", width=90, stretch=False)
        tree.column("fecha_em", anchor="center", width=115, stretch=False)
        tree.column("fecha_ven", anchor="center", width=115, stretch=False)

        ejemplos = [
            ("OC001", "FAC001", "Empresa A", "Servicio de Mantenimiento", "Pendiente", "2023-11-01", "2023-12-01"),
            ("OC002", "FAC002", "Empresa B", "Desarrollo de Software", "Aprobado", "2023-10-15", "2023-11-15"),
            ("OC003", "FAC003", "Empresa C", "Consultoría en TI", "Cancelado", "2023-09-20", "2023-10-20"),
            ("OC004", "FAC004", "Empresa D", "Instalación de Redes", "Falta Pagar", "2023-08-25", "2023-09-25"),
            ("OC005", "FAC005", "Empresa E", "Auditoría de Sistemas", "Pendiente", "2023-07-18", "2023-08-18"),
            ("OC006", "FAC006", "Empresa F", "Migración de Datos", "Aprobado", "2023-06-22", "2023-07-22"),
            ("OC007", "FAC007", "Empresa G", "Optimización de Procesos", "Pendiente", "2023-05-30", "2023-06-30"),
            ("OC008", "FAC008", "Empresa H", "Soporte Técnico", "Cancelado", "2023-04-12", "2023-05-12"),
            ("OC009", "FAC009", "Empresa I", "Desarrollo Web", "Falta Pagar", "2023-03-28", "2023-04-28"),
            ("OC010", "FAC010", "Empresa J", "Evaluación de Seguridad", "Aprobado", "2023-02-11", "2023-03-11"),
            ("OC011", "FAC011", "Empresa K", "Diseño Gráfico", "Pendiente", "2023-01-15", "2023-02-15"),
            ("OC012", "FAC012", "Empresa L", "Análisis de Datos", "Cancelado", "2022-12-05", "2023-01-05"),
            ("OC013", "FAC013", "Empresa M", "Capacitación de Personal", "Falta Pagar", "2022-11-22", "2022-12-22"),
            ("OC014", "FAC014", "Empresa N", "Optimización de Redes", "Pendiente", "2022-10-13", "2022-11-13"),
            ("OC015", "FAC015", "Empresa O", "Desarrollo de App Móvil", "Aprobado", "2022-09-07", "2022-10-07"),
            ("OC016", "FAC016", "Empresa P", "Soporte en Ciberseguridad", "Pendiente", "2022-08-20", "2022-09-20"),
            ("OC017", "FAC017", "Empresa Q", "Automatización de Procesos", "Cancelado", "2022-07-11", "2022-08-11"),
            ("OC018", "FAC018", "Empresa R", "Asesoría en Transformación Digital", "Falta Pagar", "2022-06-02", "2022-07-02"),
            ("OC019", "FAC019", "Empresa S", "Administración de Sistemas", "Pendiente", "2022-05-19", "2022-06-19"),
            ("OC020", "FAC020", "Empresa T", "Desarrollo de Contenido", "Aprobado", "2022-04-25", "2022-05-25"),
        ]
        
        ejemplos_mostrados = ejemplos[:50]
        
        for dato in ejemplos_mostrados:
            tree.insert("", "end", values=dato)
            
        scrllbar_t_fact = ttk.Scrollbar(root, orient="vertical", command=tree.yview)
        tree.configure(yscrollcommand=scrllbar_t_fact.set)
        scrllbar_t_fact.place(x=1377, y=80, height=569)

    def abrir_ventana_cotizacion(self):
        gen_cot = tk.Toplevel(root)
        generator_cot(root, gen_cot)

    def abrir_ventana_clientes(self):
        vent_clientes = tk.Toplevel(root)
        clientes(root, vent_clientes)
        
    def abrir_buscador(self):
        buscar_doc = tk.Toplevel(root)
        buscador(root, buscar_doc)
        
    def abrir_reg_cotizacion(self):
        vent_coti = tk.Toplevel(root)
        cotizaciones(root, vent_coti)
        
    def abrir_orden_compra(self):
        vent_oc = tk.Toplevel(root)
        orden_compra(root, vent_oc)
    
    def abrir_registrar_factura(self):
        reg_fact = tk.Toplevel(root)
        registrar_factura(root, reg_fact)
        
    def abrir_detalles_factura(self):
        seg_fact = tk.Toplevel(root)
        detalles_factura(root, seg_fact)


if __name__ == "__main__":
    root = tk.Tk()
    app = ventana_inicio(root)
    root.mainloop()

